import base64
import json
import mimetypes
import os
import time
import uuid
from pathlib import Path
from tempfile import NamedTemporaryFile
import secrets
import requests
from nicegui import ui, app, background_tasks, run
from generate_notes import generate_notes_from_content
from extract_content import send_msg_to_ai
from generate_word_file import generate_word_file

from logger import (
    start_file_processing,
    log_processing_success,
    log_processing_failure,
    file_logger,
)


# Import libraries for page counting
import fitz  # PyMuPDF for PDFs
from docx import Document  # python-docx for DOCX files

import hashlib
import os
from dotenv import load_dotenv

load_dotenv()  # Add this line


app.add_static_files('/assets', os.path.join(os.path.dirname(__file__), 'assets'))

# Ensure root container fills full viewport on mobile
ui.add_head_html("""
<style>
html, body, #__nicegui_root {
    height: 100%;
    margin: 0;
    padding: 0;
}
</style>
""")

# File validation constants
MAX_PAGES = 20
MAX_FILE_SIZE_MB = 70  # Additional safety check


def count_pages(file_path):
    """
    Count pages in PDF or DOCX files
    Returns: (page_count, error_message)
    """
    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            # Count PDF pages using PyMuPDF
            doc = fitz.open(file_path)
            page_count = doc.page_count
            doc.close()
            return page_count, None

        elif file_extension == '.docx':
            # Count DOCX pages (approximate based on content)
            doc = Document(file_path)

            # Method 1: Count paragraphs and estimate pages
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])

            # Rough estimation: ~10-15 paragraphs per page
            estimated_pages = max(1, paragraph_count // 12)

            # Method 2: Count characters and estimate (more accurate)
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            char_based_pages = max(1, total_chars // 3000)  # ~3000 chars per page

            # Use the higher estimate for safety
            page_count = max(estimated_pages, char_based_pages)

            return page_count, None

        else:
            return 0, "Unsupported file format"

    except Exception as e:
        return 0, f"Error reading file: {str(e)}"


def validate_file(file_path, file_size_bytes):
    """
    Validate uploaded file against size and page limits
    Returns: (is_valid, error_message, page_count)
    """
    # Check file size first (quick check)
    file_size_mb = file_size_bytes / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        return False, f"File too large ({file_size_mb:.1f}MB). Maximum size is {MAX_FILE_SIZE_MB}MB.", 0

    # Count pages
    page_count, error = count_pages(file_path)

    if error:
        return False, error, 0

    if page_count > MAX_PAGES:
        return False, f"Document has {page_count} pages. Maximum allowed is {MAX_PAGES} pages.", page_count

    return True, None, page_count



# Admin Panel


class SecureAdminAuth:
    def __init__(self):
        self.admin_password_hash = os.getenv("ADMIN_PASSWORD_HASH")
        if not self.admin_password_hash:
            print("⚠️ No admin password set! Using default: 'admin123'")
            self.admin_password_hash = self.hash_password("admin123")

    def hash_password(self, password: str) -> str:
        salt = secrets.token_hex(16)
        password_bytes = password.encode('utf-8')
        salt_bytes = salt.encode('utf-8')
        hash_obj = hashlib.sha256(salt_bytes + password_bytes)
        password_hash = hash_obj.hexdigest()
        return f"{salt}:{password_hash}"

    def verify_password(self, password: str) -> bool:
        if not self.admin_password_hash:
            return False
        try:
            stored_salt, stored_hash = self.admin_password_hash.split(':', 1)
            password_bytes = password.encode('utf-8')
            salt_bytes = stored_salt.encode('utf-8')
            hash_obj = hashlib.sha256(salt_bytes + password_bytes)
            password_hash = hash_obj.hexdigest()
            return password_hash == stored_hash
        except:
            return False


# Create the auth system
admin_auth = SecureAdminAuth()




@ui.page('/admin')
def admin_page():
    """Beautiful modern admin dashboard"""

    session = app.storage.user

    if not session.get('admin_logged_in', False):
        show_beautiful_login()
    else:
        show_beautiful_dashboard()


def show_beautiful_login():
    """Modern login page"""
    ui.add_head_html('<title>Admin Login - NotesCraft AI</title>')

    # Add custom CSS for login
    ui.add_head_html("""
    <style>
        .login-container {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
        }
        .login-card {
            backdrop-filter: blur(10px);
            background: rgba(255, 255, 255, 0.95);
            border: 1px solid rgba(255, 255, 255, 0.2);
        }
    </style>
    """)

    with ui.column().classes('login-container absolute inset-0 flex items-center justify-center'):
        with ui.card().classes('login-card w-full max-w-md p-8 shadow-2xl rounded-2xl'):
            # Logo and title
            with ui.column().classes('items-center mb-8'):
                ui.icon('admin_panel_settings').classes('text-6xl text-indigo-600 mb-4')
                ui.label('Admin Dashboard').classes('text-3xl font-bold text-gray-800')
                ui.label('NotesCraft AI Analytics').classes('text-sm text-gray-600')

            # Login form
            password_input = ui.input('Enter Admin Password', password=True).props('outlined').classes('w-full')
            error_label = ui.label('').classes('text-red-500 text-center text-sm mt-2')

            def check_password():
                session = app.storage.user
                if admin_auth.verify_password(password_input.value):  # Use secure verification
                    session['admin_logged_in'] = True
                    ui.navigate.to('/admin')
                else:
                    error_label.text = 'Incorrect password'
                    password_input.value = ''

            ui.button('Access Dashboard', on_click=check_password).props('unelevated').classes(
                'w-full mt-6 bg-indigo-600 text-white py-3 text-lg font-semibold rounded-lg hover:bg-indigo-700'
            )

            # Allow Enter key
            password_input.on('keydown.enter', check_password)

            # Footer
            ui.label('🔒 Secure admin access only').classes('text-xs text-gray-500 text-center mt-6')


def show_beautiful_dashboard():
    """Beautiful modern dashboard"""
    ui.add_head_html('<title>Analytics Dashboard - NotesCraft AI</title>')

    # Add custom CSS for dashboard
    ui.add_head_html("""
    <style>
        .dashboard-bg {
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            min-height: 100vh;
        }
        .glass-card {
            backdrop-filter: blur(10px);
            background: rgba(255, 255, 255, 0.8);
            border: 1px solid rgba(255, 255, 255, 0.2);
        }
        .status-success { background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); }
        .status-failed { background: linear-gradient(135deg, #fa709a 0%, #fee140 100%); }
        .status-processing { background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%); }
        .metric-card {
            transition: transform 0.3s ease, box-shadow 0.3s ease;
        }
        .metric-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 20px 40px rgba(0,0,0,0.1);
        }
    </style>
    """)

    with ui.column().classes('dashboard-bg w-full min-h-screen'):
        # Header Section
        with ui.row().classes('w-full p-6 items-center justify-between'):
            with ui.column():
                ui.label('📊 Analytics Dashboard').classes('text-4xl font-bold text-gray-800')
                ui.label('Real-time insights for NotesCraft AI').classes('text-lg text-gray-600')

            with ui.row().classes('gap-4'):
                ui.link('🏠 Back to App', '/').classes(
                    'px-4 py-2 bg-white rounded-lg text-indigo-600 font-medium hover:bg-indigo-50 shadow-sm'
                )

                def logout():
                    session = app.storage.user
                    session['admin_logged_in'] = False
                    ui.navigate.to('/admin')

                ui.button('🚪 Logout', on_click=logout).props('outline').classes(
                    'px-4 py-2 text-red-600 border-red-300 hover:bg-red-50'
                )

        # Main Content
        with ui.column().classes('w-full px-6 pb-6'):
            # Quick Stats Overview
            stats = file_logger.get_stats_summary()

            with ui.row().classes('w-full gap-6 mb-8'):
                # Total Processed
                with ui.card().classes('metric-card glass-card p-6 flex-1 text-center'):
                    ui.icon('description').classes('text-4xl text-blue-600 mb-2')
                    ui.label(str(stats['total_processed'])).classes('text-3xl font-bold text-gray-800')
                    ui.label('Files Processed').classes('text-sm text-gray-600 font-medium')

                # Success Rate
                success_rate = round((stats['successful'] / max(stats['total_processed'], 1)) * 100, 1)
                with ui.card().classes('metric-card glass-card p-6 flex-1 text-center'):
                    ui.icon('check_circle').classes('text-4xl text-green-600 mb-2')
                    ui.label(f'{success_rate}%').classes('text-3xl font-bold text-gray-800')
                    ui.label('Success Rate').classes('text-sm text-gray-600 font-medium')

                # Total Tokens
                total_tokens = stats['total_extraction_tokens'] + stats['total_generation_tokens']
                with ui.card().classes('metric-card glass-card p-6 flex-1 text-center'):
                    ui.icon('psychology').classes('text-4xl text-purple-600 mb-2')
                    ui.label(f'{total_tokens:,}').classes('text-3xl font-bold text-gray-800')
                    ui.label('Total Tokens').classes('text-sm text-gray-600 font-medium')

                # Avg Processing Time
                with ui.card().classes('metric-card glass-card p-6 flex-1 text-center'):
                    ui.icon('schedule').classes('text-4xl text-orange-600 mb-2')
                    ui.label(f'{stats["average_processing_time"]}s').classes('text-3xl font-bold text-gray-800')
                    ui.label('Avg Time').classes('text-sm text-gray-600 font-medium')

            # Detailed Stats Row
            with ui.row().classes('w-full gap-6 mb-8'):
                # Status Breakdown
                with ui.card().classes('glass-card p-6 flex-1'):
                    ui.label('📈 Processing Status').classes('text-xl font-bold text-gray-800 mb-4')

                    with ui.column().classes('gap-3'):
                        # Success
                        with ui.row().classes('items-center justify-between'):
                            with ui.row().classes('items-center gap-2'):
                                ui.icon('check_circle').classes('text-green-500')
                                ui.label('Successful').classes('font-medium')
                            ui.label(str(stats['successful'])).classes('font-bold text-green-600')

                        # Failed
                        with ui.row().classes('items-center justify-between'):
                            with ui.row().classes('items-center gap-2'):
                                ui.icon('error').classes('text-red-500')
                                ui.label('Failed').classes('font-medium')
                            ui.label(str(stats['failed'])).classes('font-bold text-red-600')

                        # Downloaded
                        with ui.row().classes('items-center justify-between'):
                            with ui.row().classes('items-center gap-2'):
                                ui.icon('download').classes('text-blue-500')
                                ui.label('Downloaded').classes('font-medium')
                            ui.label(str(stats['downloaded'])).classes('font-bold text-blue-600')

                # Token Breakdown
                with ui.card().classes('glass-card p-6 flex-1'):
                    ui.label('🎯 Token Usage').classes('text-xl font-bold text-gray-800 mb-4')

                    with ui.column().classes('gap-3'):
                        # Extraction tokens
                        with ui.row().classes('items-center justify-between'):
                            with ui.row().classes('items-center gap-2'):
                                ui.icon('search').classes('text-indigo-500')
                                ui.label('Extraction').classes('font-medium')
                            ui.label(f'{stats["total_extraction_tokens"]:,}').classes('font-bold text-indigo-600')

                        # Generation tokens
                        with ui.row().classes('items-center justify-between'):
                            with ui.row().classes('items-center gap-2'):
                                ui.icon('auto_fix_high').classes('text-purple-500')
                                ui.label('Generation').classes('font-medium')
                            ui.label(f'{stats["total_generation_tokens"]:,}').classes('font-bold text-purple-600')

                        # Total
                        ui.separator().classes('my-2')
                        with ui.row().classes('items-center justify-between'):
                            ui.label('Total').classes('font-bold')
                            ui.label(f'{total_tokens:,}').classes('font-bold text-gray-800 text-lg')

            # File Processing History
            logs = file_logger.read_logs()

            if not logs:
                with ui.card().classes('glass-card p-12 text-center'):
                    ui.icon('folder_open').classes('text-6xl text-gray-400 mb-4')
                    ui.label('No files processed yet').classes('text-xl text-gray-500 font-medium')
                    ui.label('Upload and process some files to see analytics here').classes('text-gray-400')
                return

            # Files Section Header
            with ui.row().classes('w-full items-center justify-between mb-6'):
                ui.label('📁 Recent File Processing').classes('text-2xl font-bold text-gray-800')
                ui.label(f'{len(logs)} files processed').classes('text-gray-600')

            # Files Grid
            recent_logs = sorted(logs, key=lambda x: x.get('start_time', ''), reverse=True)[:10]

            with ui.column().classes('w-full gap-4'):
                for log in recent_logs:
                    show_beautiful_file_card(log)


def show_beautiful_file_card(log):
    """Beautiful file processing card"""
    filename = log.get('filename', 'Unknown file')
    status = log.get('status', 'unknown')
    start_time = log.get('start_time', '')
    total_duration = log.get('total_duration', 0)
    downloaded = log.get('downloaded', False)

    # Format time
    try:
        from datetime import datetime
        dt = datetime.fromisoformat(start_time)
        formatted_time = dt.strftime('%b %d, %Y at %H:%M')
    except:
        formatted_time = start_time

    # Determine card style and status
    if status == 'success':
        if downloaded:
            card_class = 'glass-card border-l-4 border-green-500'
            status_badge = 'bg-green-100 text-green-800'
            status_icon = '✅'
            status_text = 'Completed & Downloaded'
        else:
            card_class = 'glass-card border-l-4 border-yellow-500'
            status_badge = 'bg-yellow-100 text-yellow-800'
            status_icon = '⚠️'
            status_text = 'Completed (Not Downloaded)'
    elif status == 'failed':
        card_class = 'glass-card border-l-4 border-red-500'
        status_badge = 'bg-red-100 text-red-800'
        status_icon = '❌'
        status_text = 'Processing Failed'
    else:
        card_class = 'glass-card border-l-4 border-blue-500'
        status_badge = 'bg-blue-100 text-blue-800'
        status_icon = '🔄'
        status_text = 'Processing...'

    with ui.card().classes(f'w-full p-6 {card_class} hover:shadow-lg transition-shadow'):
        with ui.row().classes('w-full items-start justify-between'):
            # Main content
            with ui.column().classes('flex-1'):
                # File header
                with ui.row().classes('items-center gap-3 mb-3'):
                    ui.icon('description').classes('text-2xl text-gray-600')
                    ui.label(filename).classes('text-xl font-bold text-gray-800')

                    # Status badge
                    with ui.row().classes(f'px-3 py-1 rounded-full {status_badge}'):
                        ui.label(f'{status_icon} {status_text}').classes('text-sm font-medium')

                # Time and basic info
                ui.label(formatted_time).classes('text-sm text-gray-500 mb-3')

                # File metrics
                pages = log.get('page_count', 0)
                size = log.get('file_size_mb', 0)

                with ui.row().classes('gap-6 mb-4'):
                    with ui.row().classes('items-center gap-1'):
                        ui.icon('description').classes('text-sm text-gray-500')
                        ui.label(f'{pages} pages').classes('text-sm text-gray-600')

                    with ui.row().classes('items-center gap-1'):
                        ui.icon('storage').classes('text-sm text-gray-500')
                        ui.label(f'{size} MB').classes('text-sm text-gray-600')

                    if total_duration:
                        with ui.row().classes('items-center gap-1'):
                            ui.icon('schedule').classes('text-sm text-gray-500')
                            ui.label(f'{total_duration}s').classes('text-sm text-gray-600')

                # Processing details (if successful)
                if status == 'success':
                    extraction = log.get('extraction', {})
                    generation = log.get('generation', {})

                    with ui.row().classes('gap-8'):
                        # Extraction info
                        ext_tokens = extraction.get('tokens', {}).get('total', 0)
                        ext_time = extraction.get('duration', 0)
                        if ext_tokens > 0:
                            with ui.column().classes('bg-indigo-50 p-3 rounded-lg'):
                                ui.label('🔍 Extraction').classes('text-xs font-bold text-indigo-600 uppercase')
                                ui.label(f'{ext_tokens:,} tokens').classes('text-sm font-medium')
                                if ext_time:
                                    ui.label(f'{ext_time}s').classes('text-xs text-gray-500')

                        # Generation info
                        gen_tokens = generation.get('tokens', {}).get('total', 0)
                        gen_time = generation.get('duration', 0)
                        if gen_tokens > 0:
                            with ui.column().classes('bg-purple-50 p-3 rounded-lg'):
                                ui.label('🛠️ Generation').classes('text-xs font-bold text-purple-600 uppercase')
                                ui.label(f'{gen_tokens:,} tokens').classes('text-sm font-medium')
                                if gen_time:
                                    ui.label(f'{gen_time}s').classes('text-xs text-gray-500')

                # Error details (if failed)
                elif status == 'failed' and log.get('error'):
                    error = log['error']
                    error_type = error.get('error_type', 'Unknown')
                    step = error.get('processing_step', 'unknown')

                    # Convert error type to readable
                    readable_errors = {
                        "API_KEY_ERROR": "API Key Missing",
                        "API_RATE_LIMIT": "Rate Limit Exceeded",
                        "API_QUOTA_EXCEEDED": "API Quota Exceeded",
                        "FILE_EXTRACTION_ERROR": "File Reading Failed",
                        "NOTES_GENERATION_ERROR": "Notes Generation Failed"
                    }
                    readable_error = readable_errors.get(error_type, error_type.replace("_", " ").title())

                    with ui.row().classes('bg-red-50 p-3 rounded-lg items-center gap-2'):
                        ui.icon('error').classes('text-red-500')
                        ui.label(f'Failed at {step}: {readable_error}').classes('text-sm text-red-700 font-medium')

            # Actions
            with ui.column().classes('gap-2'):
                # Raw data button
                def show_complete_data():
                    with ui.dialog() as dialog, ui.card().classes('w-full max-w-4xl p-6'):
                        ui.label(f'📊 Complete Processing Data').classes('text-2xl font-bold mb-4')
                        ui.label(filename).classes('text-lg text-gray-600 mb-4')

                        json_text = json.dumps(log, indent=2)
                        ui.textarea(value=json_text).classes('w-full h-96 font-mono text-sm')

                        with ui.row().classes('justify-end mt-4'):
                            ui.button('Close', on_click=dialog.close).props('outline')
                    dialog.open()

                ui.button('📊 View Details', on_click=show_complete_data).props('outline').classes(
                    'text-sm px-4 py-2 border-gray-300 text-gray-600 hover:bg-gray-50'
                )

@ui.page('/')
def main_page():
    ui.add_head_html('<link rel="icon" href="assets/favicon.ico">')
    ui.add_head_html("""
    <script src="https://unpkg.com/@lottiefiles/lottie-player@latest/dist/lottie-player.js"></script>
    """)

    session = app.storage.user
    session.uploaded_file_path = None
    session.uploaded_file_name = "Notes"

    # --- Helper Functions ---
    def report_error(Error):
        if Error:
            try:
                requests.post(
                    'https://script.google.com/macros/s/AKfycbz6Gbht0iZ4tW7lp48x3hDYCvYIDGZbOYdwnpbmyHSQjxsdZ0D0zsx7ZU84eN9n0g2T9w/exec',
                    json={"Error": Error},
                )
            except Exception as e:
                print(f"Error reporting failed: {e}")

    def handle_error_response(error_result, step_name=""):
        """
        Handle error responses from extraction or generation modules.
        Shows user-friendly message and reports technical error.
        """
        if isinstance(error_result, dict) and "error_type" in error_result:
            # This is our new error format
            user_message = error_result["user_message"]
            technical_error = error_result["technical_error"]

            # Report the technical error for debugging
            report_error(f"{step_name}: {technical_error}")

            # Show user-friendly message in UI
            error_label.text = f"⚠️ {user_message}"

            return True  # Indicates this was an error

        return False  # Not an error

    def reset_app():
        session.uploaded_file_path = None
        session.uploaded_file_name = "Notes"
        session.processing_session_id = None
        status_label.text = ""
        download_button.visible = False
        generate_button.visible = True
        reset_button.visible = False
        error_label.text = ""
        try_again_button.visible = False
        feedback_button.visible = False
        error_report_button.visible = False
        render_upload()
        ui.notify("Ready for another file!")

    async def process_with_ai():
        if not session.uploaded_file_path or not session.uploaded_file_path.exists():
            ui.notify(message='Please Upload a File', type='warning')
            return

        ui.notify('Processing may take 5-10 minutes. Mobile devices may experience connection issues.', type='info',
                  timeout=5000)

        async def background_job():
            try:
                generate_button.visible = False
                text_extraction_animation.visible = True
                status_label.text = "Extracting content from the document..."
                ui.update()

                # --- TEXT EXTRACTION ---
                try:
                    extracted_json = await run.io_bound(
                        lambda: send_msg_to_ai(session.uploaded_file_path,session.processing_session_id)
                    )

                    # Check if extraction returned an error using our new error handler
                    if handle_error_response(extracted_json, "Text Extraction"):
                        # Error was handled, show UI elements and return
                        log_processing_failure(
                            session.processing_session_id,
                            extracted_json["error_type"],
                            extracted_json["technical_error"],
                            "extraction"
                        )


                        text_extraction_animation.visible = False
                        error_report_button.visible = True
                        try_again_button.visible = True
                        status_label.text = ""
                        ui.update()
                        return

                except Exception as e:
                    # Unexpected error during extraction

                    log_processing_failure(
                        session.processing_session_id,
                        "UNEXPECTED_ERROR",
                        f"Unexpected extraction error: {str(e)}",
                        "extraction"
                    )
                    report_error(f"Unexpected Background Error: {str(e)}")
                    error_report_button.visible = True
                    text_extraction_animation.visible = False
                    try_again_button.visible = True
                    error_label.text = "⚠️ Something unexpected happened. Please try again!"
                    ui.update()
                    return

                # --- NOTES GENERATION ---
                text_extraction_animation.visible = False
                status_label.text = "🛠 Generating Notes"
                notes_generation_animation.visible = True
                ui.update()

                try:
                    # notes_generated = await run.io_bound(generate_notes_from_content, extracted_json)

                    notes_generated = await run.io_bound(
                        lambda: generate_notes_from_content(extracted_json, session.processing_session_id)
                    )
                    # Check if notes generation returned an error
                    if handle_error_response(notes_generated, "Notes Generation"):
                        # Error was handled, show UI elements and return

                        log_processing_failure(
                            session.processing_session_id,
                            notes_generated["error_type"],
                            notes_generated["technical_error"],
                            "generation"
                        )


                        notes_generation_animation.visible = False
                        error_report_button.visible = True
                        try_again_button.visible = True
                        status_label.text = ""
                        ui.update()
                        return

                    # Additional validation for empty notes
                    if not notes_generated:


                        log_processing_failure(
                            session.processing_session_id,
                            "NOTES_GENERATION_ERROR",
                            "Notes generation returned empty content",
                            "generation"
                        )

                        report_error("Notes Generation Error: Empty content returned")
                        error_label.text = "⚠️ We couldn't generate any notes from your document. Let's try again!"
                        notes_generation_animation.visible = False
                        error_report_button.visible = True
                        try_again_button.visible = True
                        status_label.text = ""
                        ui.update()
                        return

                except Exception as e:

                    log_processing_failure(
                        session.processing_session_id,
                        "UNEXPECTED_ERROR",
                        f"Unexpected generation error: {str(e)}",
                        "generation"
                    )

                    report_error(f"Notes Generation Error: {str(e)}")
                    print(f"Error: {str(e)}")
                    error_report_button.visible = True
                    notes_generation_animation.visible = False
                    status_label.text = ""
                    error_label.text = "⚠️We encountered an issue while generating your notes. Please try again."
                    try_again_button.visible = True
                    ui.update()
                    return

                # --- WORD FILE CREATION ---
                notes_generation_animation.visible = False
                status_label.text = "📕 Preparing Word file..."
                word_file_generation_animation.visible = True
                ui.update()

                try:
                    unique_name = f"{session.uploaded_file_name}_{uuid.uuid4().hex[:6]}.docx"
                    file_generated = await run.io_bound(generate_word_file, notes_generated,
                                                        file_name=unique_name.replace(' ', '_'))

                    # --- PREPARE DOWNLOAD ---
                    with open(file_generated, 'rb') as f:
                        file_content = f.read()
                    os.remove(file_generated)

                    base64_data = base64.b64encode(file_content).decode('utf-8')
                    mime_type = mimetypes.guess_type("Notes.docx")[0] or "application/octet-stream"

                    log_processing_success(session.processing_session_id)



                    def trigger_download():
                        ui.run_javascript(f"""
                                const link = document.createElement('a');
                                link.href = "data:{mime_type};base64,{base64_data}";
                                link.download = "{session.uploaded_file_name}_Notes.docx";
                                link.click();
                            """)

                        file_logger.update_download_status(session.processing_session_id)

                    time.sleep(3.5)

                    # Make sure only one listener is active
                    download_button.on('click', trigger_download, [])
                    word_file_generation_animation.visible = False
                    download_button.visible = True
                    feedback_button.visible = True
                    status_label.text = "✅ Your Notes are Ready!"
                    reset_button.visible = True
                    ui.update()


                except Exception as e:
                    # Word file creation error
                    report_error(f"Word File Creation Error: {str(e)}")
                    error_label.text = "⚠️ Almost there! Had trouble creating the Word file. Let's retry."
                    word_file_generation_animation.visible = False
                    error_report_button.visible = True
                    try_again_button.visible = True
                    status_label.text = ""
                    ui.update()
                    return

            except Exception as e:

                # --- Catch any unexpected failures ---

                log_processing_failure(
                    session.processing_session_id,
                    "WORD_FILE_ERROR",
                    f"Word file creation error: {str(e)}",
                    "word_generation"
                )

                report_error(f"System Error: {str(e)}")
                error_report_button.visible = True
                print(f"Error: {str(e)}")
                error_label.text = "⚠️ Something unexpected happened on our end. We're on it!"
                status_label.text = ""
                text_extraction_animation.visible = False
                notes_generation_animation.visible = False
                word_file_generation_animation.visible = False
                try_again_button.visible = True
                ui.update()


        background_tasks.create(background_job())

    # --- UI Layout ---
    with ui.column().classes(
            'absolute inset-0 w-full h-full overflow-x-hidden bg-gradient-to-tr '
            'from-emerald-200 via-white to-indigo-100 px-3 sm:px-6 py-6 sm:py-12 text-base sm:text-lg'):

        with ui.column().classes('w-full items-center'):
            ui.label('NotesCraft AI') \
                .classes('text-4xl md:text-4xl font-extrabold text-emerald-800 text-center')

            ui.label('Transform your PDFs into beautiful, structured study notes.') \
                .classes('text-base md:text-lg text-gray-600 text-center mb-6 px-4')

        with ui.card().classes(
                'w-full max-w-2xl sm:max-w-3xl mx-auto p-6 sm:p-8 bg-white/90 backdrop-blur-lg shadow-2xl rounded-3xl border border-gray-200 flex flex-col items-center space-y-4'):
            upload_container = ui.column().classes('w-full items-center')

            def handle_upload(e):
                # Store file details temporarily for confirmation
                temp_file_name = e.name
                temp_content = e.content.read()

                if not temp_content:
                    ui.notify("⚠️ Upload failed: empty file received.", type="warning")
                    return

                file_size = len(temp_content)
                suffix = Path(temp_file_name).suffix or ".pdf"

                # Create temporary file for validation
                with NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                    temp_file.write(temp_content)
                    temp_file.flush()
                    os.fsync(temp_file.fileno())
                    temp_file_path = Path(temp_file.name)

                # Validate file before showing confirmation
                is_valid, error_msg, page_count = validate_file(temp_file_path, file_size)

                if not is_valid:
                    # Clean up temp file and show error
                    os.unlink(temp_file_path)
                    ui.notify(f"⚠️ {error_msg}", type="negative", timeout=5000)
                    return

                def confirm_upload():
                    # User confirmed - proceed with upload
                    session.uploaded_file_name = temp_file_name
                    session.uploaded_file_path = temp_file_path

                    dialog.close()

                    session.processing_session_id = start_file_processing(
                        temp_file_name,
                        file_size / (1024 * 1024),
                        page_count
                    )

                    # Show success UI with page count
                    upload_container.clear()
                    with upload_container:
                        with ui.card().classes(
                                'w-full max-w-xl p-6 sm:p-8 rounded-2xl border border-emerald-300 '
                                'bg-emerald-50 text-center shadow-md flex flex-col items-center justify-center'
                        ):
                            ui.icon("picture_as_pdf").classes("text-red-500 text-5xl sm:text-6xl")
                            ui.label(session.uploaded_file_name).classes(
                                "text-lg sm:text-xl font-semibold text-gray-800 mt-2")
                            ui.label(f"✅ File Uploaded Successfully ({page_count} pages)").classes(
                                "text-sm sm:text-base text-gray-600 mt-1")

                def cancel_upload():
                    # User cancelled - clean up temp file
                    os.unlink(temp_file_path)
                    dialog.close()
                    ui.notify("Upload cancelled", type="info")

                # Show confirmation dialog
                with ui.dialog() as dialog, ui.card().classes('w-full max-w-md p-6'):
                    ui.label('Confirm File Upload').classes('text-xl font-bold text-gray-800 mb-4')

                    with ui.column().classes('w-full items-center space-y-3'):
                        ui.icon("description").classes("text-blue-500 text-4xl")
                        ui.label(f'File: {temp_file_name}').classes('text-lg font-medium text-gray-700')
                        ui.label(f'Pages: {page_count}').classes('text-base text-gray-600')
                        ui.label(f'Size: {file_size / 1024 / 1024:.1f} MB').classes('text-base text-gray-600')

                        ui.label('Do you want to upload this file?').classes('text-base text-gray-700 mt-4 text-center')

                        with ui.row().classes('w-full justify-center gap-4 mt-6'):
                            ui.button('Yes, Upload', on_click=confirm_upload).props(
                                'unelevated rounded color=green text-color=white').classes(
                                'px-6 py-2 font-semibold')
                            ui.button('Cancel', on_click=cancel_upload).props(
                                'unelevated rounded color=red text-color=white').classes(
                                'px-6 py-2 font-semibold')

                dialog.open()

            def render_upload():
                upload_container.clear()
                with upload_container:
                    uploader = ui.upload(label='', on_upload=handle_upload, auto_upload=True, multiple=False).props(
                        'accept=.pdf,.docx').classes('hidden')

                    with ui.card().classes(
                            'w-full max-w-xl h-48 border-2 border-dashed border-gray-300 bg-white/80 '
                            'hover:bg-emerald-50 rounded-2xl flex flex-col items-center justify-center '
                            'cursor-pointer transition-all text-center shadow-md'
                    ).on('click', lambda: uploader.run_method('pickFiles')):
                        ui.icon('cloud_upload').classes('text-5xl text-emerald-600')
                        ui.label('Click to upload your PDF or DOCX').classes('text-lg font-medium text-gray-700')
                        ui.label('or drag and drop here').classes('text-sm text-gray-500')

                    # Add page limit info below the upload area
                    ui.label(f'Maximum {MAX_PAGES} pages allowed').classes('text-xs text-gray-400 mt-2 text-center')

            render_upload()

            with ui.column().classes('w-full items-center'):
                with ui.column().classes('w-full items-center sm:flex-row sm:justify-center sm:gap-6 mt-2'):
                    text_extraction_animation = ui.html("""
                        <lottie-player src="/assets/document-search.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    text_extraction_animation.visible = False

                    notes_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_notes.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    notes_generation_animation.visible = False

                    word_file_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_word_file.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    word_file_generation_animation.visible = False

                status_label = ui.label('').classes(
                    'text-gray-900 mt-3 text-center text-base sm:text-lg font-medium'
                )

                error_label = ui.label('').classes('mt-3 text-base sm:text-lg text-rose-500 font-semibold text-center')

            download_button = ui.button('Download Notes').props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            download_button.visible = False

            feedback_url = "https://forms.gle/gPHd66XpZ1nM17si9"

            def open_feedback_form():
                js_code = f"window.open('{feedback_url}', '_blank');"
                ui.run_javascript(js_code)

            feedback_button = ui.button("📝 Give Feedback", on_click=open_feedback_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            feedback_button.visible = False

            error_report_url = "https://forms.gle/Eqjk6SS1jtmWuXGg6"

            def open_error_report_form():
                js_code = f"window.open('{error_report_url}', '_blank');"
                ui.run_javascript(js_code)

            error_report_button = ui.button("🚩 Report Error ", on_click=open_error_report_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            error_report_button.visible = False

            generate_button = ui.button('🚀 Generate Notes', on_click=process_with_ai).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')

            reset_button = ui.button('🔄 Upload Another File', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            reset_button.visible = False

            try_again_button = ui.button('🔄 Try Again', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            try_again_button.visible = False


ui.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)),
       storage_secret=os.environ.get('STORAGE_SECRET', secrets.token_hex(32)),
       title='NotesCraft AI – Smart Notes Maker')

