from docx import Document
import os
from datetime import datetime

# File path for the Word report
FILE_PATH = "NotesCraft_Token_Report.docx"

# Headers for the table
HEADERS = [
    "Run ID",
    "Date",
    "PDF Name",
    "Pages Processed",
    "Tokens Used (Input)",
    "Tokens Used (Output)",
    "Total Tokens",
    "Requests Made",
    "Cumulative Tokens",
    "Cumulative Requests",
    "Notes",
]

def create_report():
    """Create a new Word report with headers if it doesn't exist."""
    doc = Document()
    doc.add_heading("📊 NotesCraft – Token & Request Report", level=1)
    doc.add_paragraph("Testing Goal: Track tokens and requests per 10-page PDF processing.\n")

    doc.add_heading("Run Records", level=2)
    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(HEADERS):
        hdr_cells[i].text = header

    doc.save(FILE_PATH)
    print("✅ New report created:", FILE_PATH)

def get_last_totals(doc):
    """Get last cumulative totals from the table."""
    table = doc.tables[0]
    if len(table.rows) == 1:  # only header exists
        return 0, 0
    last_row = table.rows[-1].cells
    try:
        last_tokens = int(last_row[8].text)  # cumulative tokens column
        last_requests = int(last_row[9].text)  # cumulative requests column
        return last_tokens, last_requests
    except ValueError:
        return 0, 0

def get_next_run_id(doc):
    """Get the next Run ID based on the last entry in the table."""
    table = doc.tables[0]
    if len(table.rows) == 1:  # Only header exists
        return 1
    last_row = table.rows[-1].cells
    try:
        return int(last_row[0].text) + 1
    except ValueError:
        return 1

def add_record(data):
    """Add a new record (row) into the Word report."""
    doc = Document(FILE_PATH)
    table = doc.tables[0]  # first table in the document

    cells = table.add_row().cells
    for i, value in enumerate(data):
        cells[i].text = str(value)

    doc.save(FILE_PATH)
    print("✅ Record added successfully!")

def main():
    if not os.path.exists(FILE_PATH):
        create_report()

    doc = Document(FILE_PATH)

    # Auto-generate Run ID and Date
    run_id = get_next_run_id(doc)
    date = datetime.today().strftime("%d-%b-%y")  # e.g., 03-Sep-25

    # Get last cumulative totals
    last_tokens, last_requests = get_last_totals(doc)

    print("\n--- Add a new record to NotesCraft Report ---")
    pdf_name = input("PDF Name: ")
    pages = input("Pages Processed: ")
    tokens_in = int(input("Tokens Used (Input): "))
    tokens_out = int(input("Tokens Used (Output): "))
    total_tokens = tokens_in + tokens_out
    requests = int(input("Requests Made: "))
    notes = input("Notes: ")

    # Auto-calculate cumulative totals
    cumulative_tokens = last_tokens + total_tokens
    cumulative_requests = last_requests + requests

    data = [
        run_id, date, pdf_name, pages,
        tokens_in, tokens_out, total_tokens,
        requests, cumulative_tokens, cumulative_requests, notes
    ]

    add_record(data)

if __name__ == "__main__":
    main()
