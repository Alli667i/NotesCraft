import base64
import mimetypes
import os
import time
import uuid
from pathlib import Path
from tempfile import NamedTemporaryFile
import secrets
import requests
import traceback
import logging
from datetime import datetime
from nicegui import ui, app, background_tasks, run
from process_content_to_notes import generate_notes_from_content
from process_pdf_to_json import send_msg_to_ai
from process_to_word_02 import generate_word_file

# Import libraries for page counting
import fitz  # PyMuPDF for PDFs
from docx import Document  # python-docx for DOCX files

app.add_static_files('/assets', os.path.join(os.path.dirname(__file__), 'assets'))

# Configure logging for developer debugging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app_debug.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Ensure root container fills full viewport on mobile
ui.add_head_html("""
<style>
html, body, #__nicegui_root {
    height: 100%;
    margin: 0;
    padding: 0;
}
</style>
""")

# File validation constants
MAX_PAGES = 15
MAX_FILE_SIZE_MB = 60


class ErrorHandler:
    """Centralized error handling system"""

    @staticmethod
    def log_error(error_type: str, error_details: str, user_action: str = "", additional_data: dict = None):
        """Log errors for developer debugging"""
        # Get current traceback if available
        current_traceback = traceback.format_exc() if traceback.format_exc() != "NoneType: None\n" else "No traceback available"

        error_info = {
            "timestamp": datetime.now().isoformat(),
            "component": "main_app",
            "error_type": error_type,
            "error_details": error_details,
            "user_action": user_action,
            "additional_data": additional_data or {},
            "traceback": current_traceback,
            "app_version": "1.0",
            "Error": f"[MAIN_APP] {error_type}: {error_details} | Action: {user_action}"
            # This matches your original format
        }

        logger.error(f"ERROR: {error_type} | {error_details} | Action: {user_action}")
        logger.error(f"Full traceback: {current_traceback}")

        # Send to external error reporting
        try:
            response = requests.post(
                'https://script.google.com/macros/s/AKfycbz6Gbht0iZ4tW7lp48x3hDYCvYIDGZbOYdwnpbmyHSQjxsdZ0D0zsx7ZU84eN9n0g2T9w/exec',
                json=error_info,
                timeout=10
            )
            logger.info(f"Error report sent successfully. Response: {response.status_code}")
        except Exception as e:
            logger.error(f"Failed to send error report: {e}")

    @staticmethod
    def get_user_friendly_message(error_type: str) -> tuple:
        """Convert technical errors to user-friendly messages"""
        error_messages = {
            "API_KEY_MISSING": (
                "🔧 Service Temporarily Unavailable",
                "We're experiencing a configuration issue on our end. Our team has been notified and we're working to resolve this quickly. Please try again in a few minutes."
            ),
            "API_QUOTA_EXCEEDED": (
                "⏳ Service Temporarily Busy",
                "Our AI service is currently experiencing high demand. Please wait a few minutes and try again. Your document is safe and ready to process."
            ),
            "API_AUTHENTICATION_ERROR": (
                "🔧 Service Configuration Issue",
                "We're experiencing a temporary service issue. Our technical team has been automatically notified and is working on a fix. Please try again shortly."
            ),
            "EXTRACTION_FAILED": (
                "📄 Document Processing Issue",
                "We had trouble reading your document. This can happen with complex formatting or scanned documents. Please try uploading a different document or the same document in a different format."
            ),
            "NOTES_GENERATION_FAILED": (
                "🤖 AI Processing Issue",
                "Our AI had trouble generating notes from your document. This is usually temporary. Please try again, and if the issue persists, try with a different document."
            ),
            "WORD_FILE_GENERATION_FAILED": (
                "📝 File Creation Issue",
                "We successfully generated your notes but had trouble creating the download file. Please try again - your content should process faster the second time."
            ),
            "FILE_TOO_LARGE": (
                "📁 File Size Too Large",
                f"Your file is larger than our {MAX_FILE_SIZE_MB}MB limit. Please try with a smaller file or split your document into smaller sections."
            ),
            "TOO_MANY_PAGES": (
                "📖 Document Too Long",
                f"Your document has more than {MAX_PAGES} pages. Please try with a shorter document or split it into smaller sections for better results."
            ),
            "UNSUPPORTED_FORMAT": (
                "📄 Unsupported File Format",
                "We currently support PDF and DOCX files only. Please convert your document to one of these formats and try again."
            ),
            "NETWORK_ERROR": (
                "🌐 Connection Issue",
                "We're having trouble connecting to our AI service. Please check your internet connection and try again in a moment."
            ),
            "PROCESSING_TIMEOUT": (
                "⏱️ Processing Timeout",
                "Your document is taking longer than expected to process. This can happen with complex documents. Please try again with a simpler document or try again later."
            ),
            "UNKNOWN_ERROR": (
                "🚧 Unexpected Issue",
                "Something unexpected happened while processing your document. Our team has been automatically notified. Please try again, and if the issue continues, use the 'Report Error' button below."
            )
        }

        return error_messages.get(error_type, error_messages["UNKNOWN_ERROR"])


def count_pages(file_path):
    """Count pages in PDF or DOCX files with enhanced error handling"""
    try:
        file_extension = Path(file_path).suffix.lower()
        logger.info(f"Counting pages for file: {file_path} (type: {file_extension})")

        if file_extension == '.pdf':
            doc = fitz.open(file_path)
            page_count = doc.page_count
            doc.close()
            logger.info(f"PDF page count: {page_count}")
            return page_count, None

        elif file_extension == '.docx':
            doc = Document(file_path)
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            estimated_pages = max(1, paragraph_count // 12)
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            char_based_pages = max(1, total_chars // 3000)
            page_count = max(estimated_pages, char_based_pages)
            logger.info(f"DOCX estimated page count: {page_count}")
            return page_count, None

        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            return 0, "UNSUPPORTED_FORMAT"

    except Exception as e:
        ErrorHandler.log_error("PAGE_COUNT_ERROR", str(e), f"Counting pages for {file_path}")
        return 0, "UNKNOWN_ERROR"


def validate_file(file_path, file_size_bytes):
    """Validate uploaded file with enhanced error handling"""
    try:
        logger.info(f"Validating file: {file_path}, size: {file_size_bytes} bytes")

        # Check file size
        file_size_mb = file_size_bytes / (1024 * 1024)
        if file_size_mb > MAX_FILE_SIZE_MB:
            logger.warning(f"File too large: {file_size_mb:.1f}MB > {MAX_FILE_SIZE_MB}MB")
            return False, "FILE_TOO_LARGE", 0

        # Count pages
        page_count, error = count_pages(file_path)

        if error:
            return False, error, 0

        if page_count > MAX_PAGES:
            logger.warning(f"Too many pages: {page_count} > {MAX_PAGES}")
            return False, "TOO_MANY_PAGES", page_count

        logger.info(f"File validation successful: {page_count} pages, {file_size_mb:.1f}MB")
        return True, None, page_count

    except Exception as e:
        ErrorHandler.log_error("FILE_VALIDATION_ERROR", str(e), f"Validating {file_path}")
        return False, "UNKNOWN_ERROR", 0


@ui.page('/')
def main_page():
    ui.add_head_html('<link rel="icon" href="assets/favicon.ico">')
    ui.add_head_html("""
    <script src="https://unpkg.com/@lottiefiles/lottie-player@latest/dist/lottie-player.js"></script>
    """)

    session = app.storage.user
    session.uploaded_file_path = None
    session.uploaded_file_name = "Notes"

    def show_user_error(error_type: str, additional_context: str = ""):
        """Display user-friendly error messages"""
        title, message = ErrorHandler.get_user_friendly_message(error_type)

        if additional_context:
            message += f"\n\n{additional_context}"

        error_label.text = f"{title}\n{message}"
        error_label.visible = True
        try_again_button.visible = True
        error_report_button.visible = True

        # Hide all processing animations
        text_extraction_animation.visible = False
        notes_generation_animation.visible = False
        word_file_generation_animation.visible = False
        generate_button.visible = False
        status_label.text = ""

    def reset_app():
        """Reset application state"""
        try:
            logger.info("Resetting application state")
            session.uploaded_file_path = None
            session.uploaded_file_name = "Notes"
            status_label.text = ""
            download_button.visible = False
            generate_button.visible = True
            reset_button.visible = False
            error_label.text = ""
            error_label.visible = False
            try_again_button.visible = False
            feedback_button.visible = False
            error_report_button.visible = False
            render_upload()
            ui.notify("Ready for another file!")
        except Exception as e:
            ErrorHandler.log_error("RESET_ERROR", str(e), "Resetting application")

    async def process_with_ai():
        """Main processing function with comprehensive error handling"""
        if not session.uploaded_file_path or not session.uploaded_file_path.exists():
            ui.notify(message='Please Upload a File', type='warning')
            return

        async def background_job():
            try:
                logger.info(f"Starting AI processing for file: {session.uploaded_file_path}")

                generate_button.visible = False
                text_extraction_animation.visible = True
                status_label.text = "Extracting content from the document..."
                error_label.visible = False
                ui.update()

                # TEXT EXTRACTION PHASE
                try:
                    logger.info("Starting text extraction phase")
                    extracted_json = await run.io_bound(
                        lambda: send_msg_to_ai(session.uploaded_file_path)
                    )

                    # Check for extraction errors
                    if isinstance(extracted_json, str):
                        if "API" in extracted_json.upper() and "KEY" in extracted_json.upper():
                            ErrorHandler.log_error("API_KEY_MISSING", extracted_json, "Text extraction")
                            show_user_error("API_KEY_MISSING")
                        elif "QUOTA" in extracted_json.upper() or "LIMIT" in extracted_json.upper():
                            ErrorHandler.log_error("API_QUOTA_EXCEEDED", extracted_json, "Text extraction")
                            show_user_error("API_QUOTA_EXCEEDED")
                        elif "AUTHENTICATION" in extracted_json.upper() or "AUTH" in extracted_json.upper():
                            ErrorHandler.log_error("API_AUTHENTICATION_ERROR", extracted_json, "Text extraction")
                            show_user_error("API_AUTHENTICATION_ERROR")
                        elif "NETWORK" in extracted_json.upper() or "CONNECTION" in extracted_json.upper():
                            ErrorHandler.log_error("NETWORK_ERROR", extracted_json, "Text extraction")
                            show_user_error("NETWORK_ERROR")
                        else:
                            ErrorHandler.log_error("EXTRACTION_FAILED", extracted_json, "Text extraction")
                            show_user_error("EXTRACTION_FAILED")
                        return

                    if not extracted_json or len(extracted_json) == 0:
                        ErrorHandler.log_error("EXTRACTION_EMPTY", "No content extracted", "Text extraction")
                        show_user_error("EXTRACTION_FAILED", "The document appears to be empty or unreadable.")
                        return

                    logger.info("Text extraction completed successfully")

                except Exception as e:
                    error_msg = str(e)
                    if "timeout" in error_msg.lower():
                        ErrorHandler.log_error("PROCESSING_TIMEOUT", error_msg, "Text extraction")
                        show_user_error("PROCESSING_TIMEOUT")
                    elif "network" in error_msg.lower() or "connection" in error_msg.lower():
                        ErrorHandler.log_error("NETWORK_ERROR", error_msg, "Text extraction")
                        show_user_error("NETWORK_ERROR")
                    else:
                        ErrorHandler.log_error("EXTRACTION_UNEXPECTED_ERROR", error_msg, "Text extraction")
                        show_user_error("EXTRACTION_FAILED")
                    return

                # NOTES GENERATION PHASE
                text_extraction_animation.visible = False
                status_label.text = "🛠 Generating Notes"
                notes_generation_animation.visible = True
                ui.update()

                try:
                    logger.info("Starting notes generation phase")
                    notes_generated = await run.io_bound(generate_notes_from_content, extracted_json)

                    if not notes_generated:
                        ErrorHandler.log_error("NOTES_GENERATION_EMPTY", "No notes generated", "Notes generation")
                        show_user_error("NOTES_GENERATION_FAILED")
                        return

                    if isinstance(notes_generated, str) and "error" in notes_generated.lower():
                        if "api" in notes_generated.lower() and "key" in notes_generated.lower():
                            ErrorHandler.log_error("API_KEY_MISSING", notes_generated, "Notes generation")
                            show_user_error("API_KEY_MISSING")
                        elif "quota" in notes_generated.lower() or "limit" in notes_generated.lower():
                            ErrorHandler.log_error("API_QUOTA_EXCEEDED", notes_generated, "Notes generation")
                            show_user_error("API_QUOTA_EXCEEDED")
                        else:
                            ErrorHandler.log_error("NOTES_GENERATION_ERROR", notes_generated, "Notes generation")
                            show_user_error("NOTES_GENERATION_FAILED")
                        return

                    logger.info("Notes generation completed successfully")

                except Exception as e:
                    error_msg = str(e)
                    if "timeout" in error_msg.lower():
                        ErrorHandler.log_error("PROCESSING_TIMEOUT", error_msg, "Notes generation")
                        show_user_error("PROCESSING_TIMEOUT")
                    elif "quota" in error_msg.lower() or "limit" in error_msg.lower():
                        ErrorHandler.log_error("API_QUOTA_EXCEEDED", error_msg, "Notes generation")
                        show_user_error("API_QUOTA_EXCEEDED")
                    else:
                        ErrorHandler.log_error("NOTES_GENERATION_UNEXPECTED_ERROR", error_msg, "Notes generation")
                        show_user_error("NOTES_GENERATION_FAILED")
                    return

                # WORD FILE GENERATION PHASE
                notes_generation_animation.visible = False
                status_label.text = "📕 Preparing Word file..."
                word_file_generation_animation.visible = True
                ui.update()

                try:
                    logger.info("Starting Word file generation phase")
                    unique_name = f"{session.uploaded_file_name}_{uuid.uuid4().hex[:6]}.docx"
                    file_generated = await run.io_bound(generate_word_file, notes_generated,
                                                        file_name=unique_name.replace(' ', '_'))

                    if not file_generated or not os.path.exists(file_generated):
                        ErrorHandler.log_error("WORD_FILE_NOT_CREATED", "File generation returned empty path",
                                               "Word file generation")
                        show_user_error("WORD_FILE_GENERATION_FAILED")
                        return

                    # Prepare download
                    with open(file_generated, 'rb') as f:
                        file_content = f.read()
                    os.remove(file_generated)

                    base64_data = base64.b64encode(file_content).decode('utf-8')
                    mime_type = mimetypes.guess_type("Notes.docx")[0] or "application/octet-stream"

                    def trigger_download():
                        ui.run_javascript(f"""
                            const link = document.createElement('a');
                            link.href = "data:{mime_type};base64,{base64_data}";
                            link.download = "{session.uploaded_file_name}_Notes.docx";
                            link.click();
                        """)

                    time.sleep(3)

                    download_button.on('click', trigger_download, [])
                    word_file_generation_animation.visible = False
                    download_button.visible = True
                    feedback_button.visible = True
                    status_label.text = "✅ Your Notes are Ready!"
                    reset_button.visible = True
                    ui.update()

                    logger.info("Processing completed successfully")

                except Exception as e:
                    ErrorHandler.log_error("WORD_FILE_GENERATION_ERROR", str(e), "Word file generation")
                    show_user_error("WORD_FILE_GENERATION_FAILED")
                    return

            except Exception as e:
                # Catch any unexpected failures
                ErrorHandler.log_error("PROCESSING_UNEXPECTED_ERROR", str(e), "Background processing")
                show_user_error("UNKNOWN_ERROR")

        background_tasks.create(background_job())

    # UI Layout
    with ui.column().classes(
            'absolute inset-0 w-full h-full overflow-x-hidden bg-gradient-to-tr '
            'from-emerald-200 via-white to-indigo-100 px-3 sm:px-6 py-6 sm:py-12 text-base sm:text-lg'):

        with ui.column().classes('w-full items-center'):
            ui.label('NotesCraft AI') \
                .classes('text-4xl md:text-4xl font-extrabold text-emerald-800 text-center')

            ui.label('Transform your PDFs into beautiful, structured study notes.') \
                .classes('text-base md:text-lg text-gray-600 text-center mb-6 px-4')

        with ui.card().classes(
                'w-full max-w-2xl sm:max-w-3xl mx-auto p-6 sm:p-8 bg-white/90 backdrop-blur-lg shadow-2xl rounded-3xl border border-gray-200 flex flex-col items-center space-y-4'):
            upload_container = ui.column().classes('w-full items-center')

            def handle_upload(e):
                try:
                    logger.info(f"File upload initiated: {e.name}")

                    temp_file_name = e.name
                    temp_content = e.content.read()

                    if not temp_content:
                        logger.warning("Empty file uploaded")
                        ui.notify("⚠️ Upload failed: empty file received.", type="warning")
                        return

                    file_size = len(temp_content)
                    suffix = Path(temp_file_name).suffix or ".pdf"

                    # Create temporary file for validation
                    with NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                        temp_file.write(temp_content)
                        temp_file.flush()
                        os.fsync(temp_file.fileno())
                        temp_file_path = Path(temp_file.name)

                    # Validate file
                    is_valid, error_type, page_count = validate_file(temp_file_path, file_size)

                    if not is_valid:
                        os.unlink(temp_file_path)
                        title, message = ErrorHandler.get_user_friendly_message(error_type)
                        ui.notify(f"{title}: {message}", type="negative", timeout=8000)
                        ErrorHandler.log_error("FILE_VALIDATION_FAILED", f"{error_type} for file {temp_file_name}",
                                               "File upload")
                        return

                    def confirm_upload():
                        session.uploaded_file_name = temp_file_name
                        session.uploaded_file_path = temp_file_path
                        dialog.close()

                        upload_container.clear()
                        with upload_container:
                            with ui.card().classes(
                                    'w-full max-w-xl p-6 sm:p-8 rounded-2xl border border-emerald-300 '
                                    'bg-emerald-50 text-center shadow-md flex flex-col items-center justify-center'
                            ):
                                ui.icon("picture_as_pdf").classes("text-red-500 text-5xl sm:text-6xl")
                                ui.label(session.uploaded_file_name).classes(
                                    "text-lg sm:text-xl font-semibold text-gray-800 mt-2")
                                ui.label(f"✅ File Uploaded Successfully ({page_count} pages)").classes(
                                    "text-sm sm:text-base text-gray-600 mt-1")

                        logger.info(f"File upload confirmed: {temp_file_name}")

                    def cancel_upload():
                        os.unlink(temp_file_path)
                        dialog.close()
                        ui.notify("Upload cancelled", type="info")
                        logger.info(f"File upload cancelled: {temp_file_name}")

                    # Show confirmation dialog
                    with ui.dialog() as dialog, ui.card().classes('w-full max-w-md p-6'):
                        ui.label('Confirm File Upload').classes('text-xl font-bold text-gray-800 mb-4')

                        with ui.column().classes('w-full items-center space-y-3'):
                            ui.icon("description").classes("text-blue-500 text-4xl")
                            ui.label(f'File: {temp_file_name}').classes('text-lg font-medium text-gray-700')
                            ui.label(f'Pages: {page_count}').classes('text-base text-gray-600')
                            ui.label(f'Size: {file_size / 1024 / 1024:.1f} MB').classes('text-base text-gray-600')

                            ui.label('Do you want to upload this file?').classes(
                                'text-base text-gray-700 mt-4 text-center')

                            with ui.row().classes('w-full justify-center gap-4 mt-6'):
                                ui.button('Yes, Upload', on_click=confirm_upload).props(
                                    'unelevated rounded color=green text-color=white').classes(
                                    'px-6 py-2 font-semibold')
                                ui.button('Cancel', on_click=cancel_upload).props(
                                    'unelevated rounded color=red text-color=white').classes(
                                    'px-6 py-2 font-semibold')

                    dialog.open()

                except Exception as e:
                    ErrorHandler.log_error("UPLOAD_HANDLER_ERROR", str(e), "File upload handling")
                    ui.notify("Upload failed due to an unexpected error. Please try again.", type="negative")

            def render_upload():
                upload_container.clear()
                with upload_container:
                    uploader = ui.upload(label='', on_upload=handle_upload, auto_upload=True, multiple=False).props(
                        'accept=.pdf,.docx').classes('hidden')

                    with ui.card().classes(
                            'w-full max-w-xl h-48 border-2 border-dashed border-gray-300 bg-white/80 '
                            'hover:bg-emerald-50 rounded-2xl flex flex-col items-center justify-center '
                            'cursor-pointer transition-all text-center shadow-md'
                    ).on('click', lambda: uploader.run_method('pickFiles')):
                        ui.icon('cloud_upload').classes('text-5xl text-emerald-600')
                        ui.label('Click to upload your PDF or DOCX').classes('text-lg font-medium text-gray-700')
                        ui.label('or drag and drop here').classes('text-sm text-gray-500')

                    ui.label(f'Maximum {MAX_PAGES} pages allowed').classes('text-xs text-gray-400 mt-2 text-center')

            render_upload()

            # Processing animations and status
            with ui.column().classes('w-full items-center'):
                with ui.column().classes('w-full items-center sm:flex-row sm:justify-center sm:gap-6 mt-2'):
                    text_extraction_animation = ui.html("""
                        <lottie-player src="/assets/document-search.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    text_extraction_animation.visible = False

                    notes_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_notes.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    notes_generation_animation.visible = False

                    word_file_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_word_file.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    word_file_generation_animation.visible = False

                status_label = ui.label('').classes(
                    'text-gray-900 mt-3 text-center text-base sm:text-lg font-medium'
                )

                error_label = ui.label('').classes(
                    'mt-3 text-base sm:text-lg text-rose-500 font-semibold text-center whitespace-pre-line')
                error_label.visible = False

            # Action buttons
            download_button = ui.button('Download Notes').props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            download_button.visible = False

            feedback_url = "https://forms.gle/gPHd66XpZ1nM17si9"

            def open_feedback_form():
                js_code = f"window.open('{feedback_url}', '_blank');"
                ui.run_javascript(js_code)

            feedback_button = ui.button("📝 Give Feedback", on_click=open_feedback_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            feedback_button.visible = False

            error_report_url = "https://forms.gle/Eqjk6SS1jtmWuXGg6"

            def open_error_report_form():
                js_code = f"window.open('{error_report_url}', '_blank');"
                ui.run_javascript(js_code)

            error_report_button = ui.button("🚩 Report Error", on_click=open_error_report_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            error_report_button.visible = False

            generate_button = ui.button('🚀 Generate Notes', on_click=process_with_ai).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')

            reset_button = ui.button('🔄 Upload Another File', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            reset_button.visible = False

            try_again_button = ui.button('🔄 Try Again', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            try_again_button.visible = False



ui.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)),
       storage_secret=os.environ.get('STORAGE_SECRET', secrets.token_hex(32)),
       title='NotesCraft AI – Smart Notes Maker')