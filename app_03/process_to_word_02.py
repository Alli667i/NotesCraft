from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph_with_spacing(doc, text, font_size=12, bold=False, align='left', space_before=12, space_after=6):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'

    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)

def generate_word_file(content, file_name):
    doc = Document()  # <-- moved inside function

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for item in content:
        try:
            item_type = item.get('type')
            text = item.get('text', '')

            if item_type == 'heading':
                add_paragraph_with_spacing(doc, text, font_size=16, bold=True, align='center', space_before=48, space_after=36)

            elif item_type == 'subheading':
                add_paragraph_with_spacing(doc, text, font_size=14, bold=True, align='left', space_before=36, space_after=24)

            elif item_type == 'paragraph':
                add_paragraph_with_spacing(doc, text, font_size=12, bold=False, align='justify', space_before=24, space_after=24)

            elif item_type == 'bullet':
                bullet_para = doc.add_paragraph()
                run = bullet_para.add_run(f"• {text}")
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                bullet_para.paragraph_format.space_before = Pt(24)
                bullet_para.paragraph_format.space_after = Pt(28)
                bullet_para.paragraph_format.left_indent = Pt(18)
                bullet_para.paragraph_format.line_spacing = Pt(16)



        except Exception as e:
            print(f"Skipped invalid item: {item} due to error: {e}")

    file_path = f"{file_name}.docx"
    doc.save(file_path)
    return file_path
