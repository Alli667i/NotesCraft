import base64
import mimetypes
import os
import time
import uuid
from pathlib import Path
from tempfile import NamedTemporaryFile
import secrets
import requests
import traceback
import logging
from datetime import datetime
from nicegui import ui, app, background_tasks, run
from process_content_to_notes import generate_notes_from_content
from process_pdf_to_json import send_msg_to_ai
from process_to_word_02 import generate_word_file
from logging_system import production_logger, log_info, log_warning, log_error
import json
# Import libraries for page counting
import fitz  # PyMuPDF for PDFs
from docx import Document  # python-docx for DOCX files
import hashlib
import secrets




app.add_static_files('/assets', os.path.join(os.path.dirname(__file__), 'assets'))

# Configure logging for developer debugging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app_debug.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Ensure root container fills full viewport on mobile
ui.add_head_html("""
<style>
html, body, #__nicegui_root {
    height: 100%;
    margin: 0;
    padding: 0;
}
</style>
<script>
// Better mobile connection handling - NO PAGE RELOAD
let wasHidden = false;
let uploadInProgress = false;

// Track when user leaves app (to file manager)
document.addEventListener('visibilitychange', function() {
    if (document.hidden) {
        wasHidden = true;
        // User left app - possibly going to file manager
        console.log('User left app - preserving state');
    } else if (wasHidden) {
        // User returned - don't reload, just reconnect gracefully
        console.log('User returned - maintaining connection');
        wasHidden = false;

        // Give a moment for any pending uploads to complete
        setTimeout(() => {
            // Try to reconnect WebSocket without page reload
            if (window.socket && window.socket.readyState !== 1) {
                console.log('Attempting graceful reconnection...');
                // Instead of reload, try to reconnect the socket
                try {
                    window.socket.close();
                    // NiceGUI will automatically create new connection
                } catch (e) {
                    console.log('Socket cleanup completed');
                }
            }
        }, 500);
    }
});

// Prevent browser from killing the app during file selection
document.addEventListener('click', function(e) {
    // If user clicks on file upload area
    if (e.target.closest('.q-uploader') || 
        e.target.closest('[role="button"]') ||
        e.target.matches('input[type="file"]')) {

        uploadInProgress = true;
        console.log('File selection started - preserving app state');

        // Keep app alive during file selection
        const keepAlive = setInterval(() => {
            if (!document.hidden && uploadInProgress) {
                // Send tiny ping to keep connection alive
                fetch(window.location.href, { 
                    method: 'HEAD',
                    cache: 'no-cache'
                }).catch(() => {});
            } else {
                clearInterval(keepAlive);
            }
        }, 2000);

        // Stop keep-alive after 30 seconds
        setTimeout(() => {
            uploadInProgress = false;
            clearInterval(keepAlive);
        }, 30000);
    }
});

// Handle successful file selection
document.addEventListener('change', function(e) {
    if (e.target.type === 'file') {
        uploadInProgress = false;
        console.log('File selected successfully');
    }
});

// Prevent page unload during critical operations
window.addEventListener('beforeunload', function(e) {
    if (uploadInProgress) {
        e.preventDefault();
        e.returnValue = 'File selection in progress...';
        return 'File selection in progress...';
    }
});
</script>
""")

# File validation constants
MAX_PAGES = 15
MAX_FILE_SIZE_MB = 60


class SecureAuth:
    """Secure authentication system for log viewer"""

    def __init__(self):
        # Get password from environment variable or generate one
        self.admin_password = os.environ.get('LOG_ADMIN_PASSWORD')

        if not self.admin_password:
            # Generate a random password if none set
            self.admin_password = secrets.token_urlsafe(12)
            print(f"\n🔐 LOG VIEWER PASSWORD: {self.admin_password}")
            print("💡 Set LOG_ADMIN_PASSWORD environment variable to use a custom password\n")

        # Hash the password for comparison
        self.password_hash = hashlib.sha256(self.admin_password.encode()).hexdigest()

        # Session management
        self.active_sessions = {}
        self.session_timeout = 3600  # 1 hour

        # Rate limiting
        self.failed_attempts = {}  # Track failed login attempts by IP
        self.max_attempts = 5  # Max failed attempts before lockout
        self.lockout_duration = 300  # 5 minutes lockout

    def verify_password(self, input_password: str) -> bool:
        """Verify password without storing it"""
        if not input_password:
            return False

        input_hash = hashlib.sha256(input_password.encode()).hexdigest()
        return secrets.compare_digest(self.password_hash, input_hash)

    def create_session(self, session_id: str) -> str:
        """Create a secure session token"""
        token = secrets.token_urlsafe(32)
        self.active_sessions[session_id] = {
            'token': token,
            'created': time.time(),
            'last_activity': time.time()
        }
        return token

    def verify_session(self, session_id: str, token: str) -> bool:
        """Verify session is valid and not expired"""
        if session_id not in self.active_sessions:
            return False

        session = self.active_sessions[session_id]

        # Check if session expired
        if time.time() - session['last_activity'] > self.session_timeout:
            self.revoke_session(session_id)
            return False

        # Check token
        if not secrets.compare_digest(session.get('token', ''), token):
            return False

        # Update last activity
        session['last_activity'] = time.time()
        return True

    def revoke_session(self, session_id: str):
        """Revoke a session"""
        if session_id in self.active_sessions:
            del self.active_sessions[session_id]

    def cleanup_expired_sessions(self):
        """Remove expired sessions"""
        current_time = time.time()
        expired_sessions = [
            session_id for session_id, session in self.active_sessions.items()
            if current_time - session['last_activity'] > self.session_timeout
        ]

        for session_id in expired_sessions:
            del self.active_sessions[session_id]

    def is_ip_locked(self, ip_address: str) -> bool:
        """Check if IP is locked due to too many failed attempts"""
        if ip_address not in self.failed_attempts:
            return False

        attempts = self.failed_attempts[ip_address]
        if attempts['count'] >= self.max_attempts:
            # Check if lockout period has passed
            time_since_last = time.time() - attempts['last_attempt']
            if time_since_last < self.lockout_duration:
                return True
            else:
                # Reset attempts after lockout period
                del self.failed_attempts[ip_address]
                return False
        return False

    def record_failed_attempt(self, ip_address: str):
        """Record a failed login attempt"""
        current_time = time.time()

        if ip_address not in self.failed_attempts:
            self.failed_attempts[ip_address] = {'count': 0, 'last_attempt': current_time}

        self.failed_attempts[ip_address]['count'] += 1
        self.failed_attempts[ip_address]['last_attempt'] = current_time

    def clear_failed_attempts(self, ip_address: str):
        """Clear failed attempts for successful login"""
        if ip_address in self.failed_attempts:
            del self.failed_attempts[ip_address]


# 3. INITIALIZE THE AUTH SYSTEM (add this line after your other initializations)
secure_auth = SecureAuth()



def add_security_headers():
    """Add security headers for production"""
    ui.add_head_html("""
    <meta http-equiv="X-Content-Type-Options" content="nosniff">
    <meta http-equiv="X-Frame-Options" content="DENY">
    <meta http-equiv="X-XSS-Protection" content="1; mode=block">
    <meta http-equiv="Referrer-Policy" content="strict-origin-when-cross-origin">
    <meta http-equiv="Permissions-Policy" content="geolocation=(), microphone=(), camera=()">
    """)


class ErrorHandler:
    """Centralized error handling system with production logging"""

    @staticmethod
    def log_error(error_type: str, error_details: str, user_action: str = "", additional_data: dict = None):
        """Log errors using production logging system"""

        # Log to our production system
        log_error(
            component="main_app",
            message=f"{error_type}: {error_details}",
            user_action=user_action,
            additional_data=additional_data or {},
            error_details=error_details
        )

        # Still send critical errors to Google Apps Script (but async and smaller payload)
        try:
            # Only send critical info to avoid large payloads
            critical_info = {
                "timestamp": datetime.now().isoformat(),
                "error_type": error_type,
                "error_summary": error_details[:200],  # Truncate long errors
                "user_action": user_action,
                "app_version": "1.0"
            }

            # Send in background to avoid blocking
            import threading
            def send_critical_error():
                try:
                    import requests
                    requests.post(
                        'https://script.google.com/macros/s/AKfycbz6Gbht0iZ4tW7lp48x3hDYCvYIDGZbOYdwnpbmyHSQjxsdZ0D0zsx7ZU84eN9n0g2T9w/exec',
                        json=critical_info,
                        timeout=5  # Shorter timeout
                    )
                except:
                    pass  # Don't let external logging break the app

            threading.Thread(target=send_critical_error, daemon=True).start()

        except Exception:
            pass  # Don't let logging errors break the app

    @staticmethod
    def get_user_friendly_message(error_type: str) -> tuple:
        """Convert technical errors to user-friendly messages"""
        error_messages = {
            "API_KEY_MISSING": (
                "🔧 Service Temporarily Unavailable",
                "We're experiencing a configuration issue on our end. Our team has been notified and we're working to resolve this quickly. Please try again in a few minutes."
            ),
            "API_QUOTA_EXCEEDED": (
                "⏳ Service Temporarily Busy",
                "Our AI service is currently experiencing high demand. Please wait a few minutes and try again. Your document is safe and ready to process."
            ),
            "API_AUTHENTICATION_ERROR": (
                "🔧 Service Configuration Issue",
                "We're experiencing a temporary service issue. Our technical team has been automatically notified and is working on a fix. Please try again shortly."
            ),
            "EXTRACTION_FAILED": (
                "📄 Document Processing Issue",
                "We had trouble reading your document. This can happen with complex formatting or scanned documents. Please try uploading a different document or the same document in a different format."
            ),
            "NOTES_GENERATION_FAILED": (
                "🤖 AI Processing Issue",
                "Our AI had trouble generating notes from your document. This is usually temporary. Please try again, and if the issue persists, try with a different document."
            ),
            "WORD_FILE_GENERATION_FAILED": (
                "📝 File Creation Issue",
                "We successfully generated your notes but had trouble creating the download file. Please try again - your content should process faster the second time."
            ),
            "FILE_TOO_LARGE": (
                "📁 File Size Too Large",
                f"Your file is larger than our {MAX_FILE_SIZE_MB}MB limit. Please try with a smaller file or split your document into smaller sections."
            ),
            "TOO_MANY_PAGES": (
                "📖 Document Too Long",
                f"Your document has more than {MAX_PAGES} pages. Please try with a shorter document or split it into smaller sections for better results."
            ),
            "UNSUPPORTED_FORMAT": (
                "📄 Unsupported File Format",
                "We currently support PDF and DOCX files only. Please convert your document to one of these formats and try again."
            ),
            "NETWORK_ERROR": (
                "🌐 Connection Issue",
                "We're having trouble connecting to our AI service. Please check your internet connection and try again in a moment."
            ),
            "PROCESSING_TIMEOUT": (
                "⏱️ Processing Timeout",
                "Your document is taking longer than expected to process. This can happen with complex documents. Please try again with a simpler document or try again later."
            ),
            "UNKNOWN_ERROR": (
                "🚧 Unexpected Issue",
                "Something unexpected happened while processing your document. Our team has been automatically notified. Please try again, and if the issue continues, use the 'Report Error' button below."
            )
        }

        return error_messages.get(error_type, error_messages["UNKNOWN_ERROR"])


def count_pages(file_path):
    """Count pages in PDF or DOCX files with enhanced error handling"""
    try:
        file_extension = Path(file_path).suffix.lower()
        logger.info(f"Counting pages for file: {file_path} (type: {file_extension})")

        if file_extension == '.pdf':
            doc = fitz.open(file_path)
            page_count = doc.page_count
            doc.close()
            logger.info(f"PDF page count: {page_count}")
            return page_count, None

        elif file_extension == '.docx':
            doc = Document(file_path)
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            estimated_pages = max(1, paragraph_count // 12)
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            char_based_pages = max(1, total_chars // 3000)
            page_count = max(estimated_pages, char_based_pages)
            logger.info(f"DOCX estimated page count: {page_count}")
            return page_count, None

        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            return 0, "UNSUPPORTED_FORMAT"

    except Exception as e:
        ErrorHandler.log_error("PAGE_COUNT_ERROR", str(e), f"Counting pages for {file_path}")
        return 0, "UNKNOWN_ERROR"


def validate_file(file_path, file_size_bytes):
    """Validate uploaded file with enhanced error handling"""
    try:
        logger.info(f"Validating file: {file_path}, size: {file_size_bytes} bytes")

        # Check file size
        file_size_mb = file_size_bytes / (1024 * 1024)
        if file_size_mb > MAX_FILE_SIZE_MB:
            logger.warning(f"File too large: {file_size_mb:.1f}MB > {MAX_FILE_SIZE_MB}MB")
            return False, "FILE_TOO_LARGE", 0

        # Count pages
        page_count, error = count_pages(file_path)

        if error:
            return False, error, 0

        if page_count > MAX_PAGES:
            logger.warning(f"Too many pages: {page_count} > {MAX_PAGES}")
            return False, "TOO_MANY_PAGES", page_count

        logger.info(f"File validation successful: {page_count} pages, {file_size_mb:.1f}MB")
        return True, None, page_count

    except Exception as e:
        ErrorHandler.log_error("FILE_VALIDATION_ERROR", str(e), f"Validating {file_path}")
        return False, "UNKNOWN_ERROR", 0


def keep_session_alive():
    """Keep the session alive during file selection"""
    try:
        # This is a lightweight ping to maintain session
        ui.run_javascript("console.log('Session maintained');")
    except:
        pass


@ui.page('/logs')
def logs_page():
    """Secure web interface for viewing logs"""

    # Clean up expired sessions
    secure_auth.cleanup_expired_sessions()

    session = app.storage.user
    session_id = session.get('session_id', str(secrets.token_urlsafe(16)))
    session['session_id'] = session_id

    # Check if user is already authenticated
    stored_token = session.get('log_auth_token')
    if stored_token and secure_auth.verify_session(session_id, stored_token):
        # User is authenticated, show logs interface
        show_logs_interface()
        return

    # Show login form
    show_login_form(session_id)


def show_login_form(session_id: str):
    """Display secure login form with rate limiting"""

    ui.add_head_html('<title>NotesCraft AI - Admin Login</title>')

    # Get client IP for rate limiting
    client_ip = "127.0.0.1"  # Default fallback

    # Check if IP is locked FIRST - and RETURN early if locked
    if secure_auth.is_ip_locked(client_ip):
        with ui.column().classes('w-full h-screen items-center justify-center bg-red-50'):
            with ui.card().classes('w-full max-w-md p-8 shadow-xl border-2 border-red-200'):
                ui.icon('block').classes('text-6xl text-red-600 mb-4 mx-auto')
                ui.label('Access Temporarily Blocked').classes('text-xl font-bold text-red-800 text-center mb-2')
                ui.label('Too many failed login attempts. Please try again in 5 minutes.').classes(
                    'text-sm text-red-600 text-center')

                # Add countdown timer
                remaining_time = int(secure_auth.lockout_duration - (
                            time.time() - secure_auth.failed_attempts[client_ip]['last_attempt']))
                if remaining_time > 0:
                    ui.label(f'Time remaining: {remaining_time // 60}:{remaining_time % 60:02d}').classes(
                        'text-sm text-red-500 mt-2')

                # Add back to home button
                ui.button('🏠 Back to Home', on_click=lambda: ui.navigate.to('/')).props('color=primary').classes('mt-4')

        return  # CRITICAL: Return here to prevent showing login form

    # Only show login form if NOT locked
    with ui.column().classes(
            'w-full h-screen items-center justify-center bg-gradient-to-br from-blue-50 to-indigo-100'):
        with ui.card().classes('w-full max-w-md p-8 shadow-xl'):

            # Header
            with ui.column().classes('w-full items-center mb-6'):
                ui.icon('admin_panel_settings').classes('text-6xl text-indigo-600 mb-2')
                ui.label('Admin Access').classes('text-2xl font-bold text-gray-800')
                ui.label('Enter password to view application logs').classes('text-sm text-gray-600 text-center')

            # Show remaining attempts warning
            if client_ip in secure_auth.failed_attempts:
                attempts_left = secure_auth.max_attempts - secure_auth.failed_attempts[client_ip]['count']
                if attempts_left > 0:
                    ui.label(f'⚠️ {attempts_left} attempts remaining').classes(
                        'text-xs text-orange-600 text-center mb-2')

            # Login form
            password_input = ui.input('Password', password=True).classes('w-full mb-4').props('outlined')
            status_label = ui.label('').classes('text-sm text-center mb-4')
            status_label.visible = False

            def attempt_login():
                password = password_input.value

                # Double-check if IP got locked during this session
                if secure_auth.is_ip_locked(client_ip):
                    status_label.text = '🚫 Access blocked. Please refresh the page.'
                    status_label.classes('text-red-600')
                    status_label.visible = True
                    password_input.visible = False  # Hide password input
                    return

                if not password:
                    status_label.text = '⚠️ Please enter a password'
                    status_label.classes('text-yellow-600')
                    status_label.visible = True
                    return

                if secure_auth.verify_password(password):
                    # Clear failed attempts on successful login
                    secure_auth.clear_failed_attempts(client_ip)

                    # Create secure session
                    session = app.storage.user
                    token = secure_auth.create_session(session_id)
                    session['log_auth_token'] = token

                    # Log successful access
                    log_info("auth", f"Admin login successful from {client_ip}",
                             user_action="admin_login",
                             additional_data={"session_id": session_id[:8] + "...", "ip": client_ip})

                    status_label.text = '✅ Login successful! Redirecting...'
                    status_label.classes('text-green-600')
                    status_label.visible = True

                    # Redirect to logs
                    ui.run_javascript('setTimeout(() => window.location.reload(), 1000);')

                else:
                    # Record failed attempt
                    secure_auth.record_failed_attempt(client_ip)

                    # Log failed attempt
                    log_warning("auth", f"Failed admin login attempt from {client_ip}",
                                user_action="admin_login_failed",
                                additional_data={"session_id": session_id[:8] + "...", "ip": client_ip})

                    # Check if this attempt caused a lockout
                    if secure_auth.is_ip_locked(client_ip):
                        status_label.text = '🚫 Too many failed attempts. Access blocked for 5 minutes. Please refresh the page.'
                        status_label.classes('text-red-600')
                        status_label.visible = True
                        password_input.visible = False  # Hide password input
                        # Automatically refresh page after 2 seconds
                        ui.run_javascript('setTimeout(() => window.location.reload(), 2000);')
                        return

                    # Show remaining attempts
                    attempts_left = secure_auth.max_attempts - secure_auth.failed_attempts[client_ip]['count']
                    status_label.text = f'❌ Incorrect password. {attempts_left} attempts remaining.'
                    status_label.classes('text-red-600')
                    status_label.visible = True
                    password_input.value = ''

            # Login button
            login_button = ui.button('🔓 Access Logs', on_click=attempt_login).props('color=primary unelevated').classes(
                'w-full py-3 text-lg font-semibold')

            # Handle Enter key
            password_input.on('keydown.enter', attempt_login)

            # Security info
            with ui.expansion('🔧 Session Info', icon='info').classes('w-full mt-4 text-xs'):
                ui.label(f'Session ID: {session_id[:8]}...').classes('text-xs font-mono')
                ui.label('Session expires after 1 hour of inactivity').classes('text-xs text-gray-500')
                ui.label('Maximum 5 login attempts before 5-minute lockout').classes('text-xs text-gray-500')




def show_logs_interface():
    """Display the improved, organized logs interface"""

    ui.add_head_html('<title>NotesCraft AI - Logs Dashboard</title>')

    with ui.column().classes('w-full max-w-7xl mx-auto p-4 space-y-6'):

        # Header with logout
        with ui.row().classes('w-full items-center justify-between bg-white p-4 rounded-lg shadow-sm border'):
            with ui.column():
                ui.label('📊 Logs Dashboard').classes('text-2xl font-bold text-gray-800')
                ui.label('Monitor your application health and user activity').classes('text-sm text-gray-600')

            with ui.row().classes('gap-2'):
                def logout():
                    session = app.storage.user
                    session_id = session.get('session_id')
                    if session_id:
                        secure_auth.revoke_session(session_id)

                    session.pop('log_auth_token', None)
                    session.pop('session_id', None)

                    log_info("auth", "Admin logout", user_action="admin_logout")
                    ui.navigate.to('/logs')

                ui.button('🔒 Logout', on_click=logout).props('color=secondary outlined')
                ui.button('🏠 Back to App', on_click=lambda: ui.navigate.to('/')).props('color=primary')

        # Quick Stats Cards
        stats_container = ui.row().classes('w-full gap-4')

        def refresh_stats():
            stats = production_logger.get_log_stats()
            stats_container.clear()

            with stats_container:
                # Total Logs Today
                today_count = stats.get('todays_log_count', 0)
                with ui.card().classes('flex-1 p-4 bg-blue-50 border border-blue-200'):
                    with ui.row().classes('items-center gap-3'):
                        ui.icon('today').classes('text-3xl text-blue-600')
                        with ui.column():
                            ui.label(str(today_count)).classes('text-2xl font-bold text-blue-800')
                            ui.label('Today\'s Logs').classes('text-sm text-blue-600')

                # Error Count
                error_count = stats.get('level_breakdown', {}).get('ERROR', 0)
                error_color = 'red' if error_count > 0 else 'green'
                with ui.card().classes(f'flex-1 p-4 bg-{error_color}-50 border border-{error_color}-200'):
                    with ui.row().classes('items-center gap-3'):
                        error_icon = 'error' if error_count > 0 else 'check_circle'
                        ui.icon(error_icon).classes(f'text-3xl text-{error_color}-600')
                        with ui.column():
                            ui.label(str(error_count)).classes(f'text-2xl font-bold text-{error_color}-800')
                            ui.label('Errors Today').classes(f'text-sm text-{error_color}-600')

                # Storage Usage
                size_mb = stats.get('total_size_mb', 0)
                with ui.card().classes('flex-1 p-4 bg-purple-50 border border-purple-200'):
                    with ui.row().classes('items-center gap-3'):
                        ui.icon('storage').classes('text-3xl text-purple-600')
                        with ui.column():
                            ui.label(f'{size_mb} MB').classes('text-2xl font-bold text-purple-800')
                            ui.label('Storage Used').classes('text-sm text-purple-600')

        refresh_stats()

        # Tabbed Interface for Different Log Types
        with ui.card().classes('w-full'):
            with ui.tabs().classes('w-full') as tabs:
                tab_all = ui.tab('🌟 Recent Activity')
                tab_errors = ui.tab('🚨 Errors Only')
                tab_user = ui.tab('👤 User Actions')
                tab_auth = ui.tab('🔐 Authentication')
                tab_system = ui.tab('⚙️ System Events')

            with ui.tab_panels(tabs, value=tab_all).classes('w-full'):

                # Recent Activity Tab
                with ui.tab_panel(tab_all):
                    ui.label('🌟 Recent Activity').classes('text-lg font-bold mb-4')
                    recent_container = ui.column().classes('w-full space-y-2')

                    def load_recent_activity():
                        recent_container.clear()
                        logs = production_logger.get_logs(limit=20)

                        if not logs:
                            with recent_container:
                                ui.label('No recent activity found.').classes('text-gray-500 text-center p-8')
                            return

                        with recent_container:
                            for log in logs:
                                display_clean_log_card(log)

                    load_recent_activity()

                # Errors Only Tab
                with ui.tab_panel(tab_errors):
                    ui.label('🚨 Application Errors').classes('text-lg font-bold mb-4')
                    errors_container = ui.column().classes('w-full space-y-2')

                    def load_errors():
                        errors_container.clear()
                        logs = production_logger.get_logs(level='ERROR', limit=50)

                        if not logs:
                            with errors_container:
                                with ui.card().classes('w-full p-8 bg-green-50 border border-green-200 text-center'):
                                    ui.icon('check_circle').classes('text-4xl text-green-600 mb-2')
                                    ui.label('🎉 No errors found!').classes('text-lg font-semibold text-green-800')
                                    ui.label('Your application is running smoothly.').classes('text-green-600')
                            return

                        with errors_container:
                            for log in logs:
                                display_error_card(log)

                    load_errors()

                # User Actions Tab
                with ui.tab_panel(tab_user):
                    ui.label('👤 User Interactions').classes('text-lg font-bold mb-4')
                    user_container = ui.column().classes('w-full space-y-2')

                    def load_user_actions():
                        user_container.clear()
                        logs = production_logger.get_logs(component='file_upload', limit=30)

                        if not logs:
                            with user_container:
                                ui.label('No user activity found.').classes('text-gray-500 text-center p-8')
                            return

                        with user_container:
                            for log in logs:
                                display_user_action_card(log)

                    load_user_actions()

                # Authentication Tab
                with ui.tab_panel(tab_auth):
                    ui.label('🔐 Login & Security Events').classes('text-lg font-bold mb-4')
                    auth_container = ui.column().classes('w-full space-y-2')

                    def load_auth_logs():
                        auth_container.clear()
                        logs = production_logger.get_logs(component='auth', limit=50)

                        if not logs:
                            with auth_container:
                                ui.label('No authentication events found.').classes('text-gray-500 text-center p-8')
                            return

                        with auth_container:
                            for log in logs:
                                display_auth_card(log)

                    load_auth_logs()

                # System Events Tab
                with ui.tab_panel(tab_system):
                    ui.label('⚙️ System & Processing Events').classes('text-lg font-bold mb-4')

                    # Date filter for system events
                    with ui.row().classes('w-full gap-4 items-end mb-4'):
                        date_input = ui.input('Date (YYYY-MM-DD)', value=datetime.now().strftime('%Y-%m-%d')).classes(
                            'flex-1')

                        def apply_date_filter():
                            load_system_events(date_input.value)

                        ui.button('🔍 Filter by Date', on_click=apply_date_filter).props('color=primary')
                        ui.button('📅 Today',
                                  on_click=lambda: [setattr(date_input, 'value', datetime.now().strftime('%Y-%m-%d')),
                                                    load_system_events()]).props('color=secondary')

                    system_container = ui.column().classes('w-full space-y-2')

                    def load_system_events(filter_date=None):
                        system_container.clear()
                        logs = production_logger.get_logs(
                            date=filter_date or datetime.now().strftime('%Y-%m-%d'),
                            limit=100
                        )

                        # Filter out auth logs for cleaner system view
                        system_logs = [log for log in logs if log.get('component') != 'auth']

                        if not system_logs:
                            with system_container:
                                ui.label('No system events found for this date.').classes(
                                    'text-gray-500 text-center p-8')
                            return

                        with system_container:
                            for log in system_logs:
                                display_system_card(log)

                    load_system_events()

        # Refresh button at the bottom
        with ui.row().classes('w-full justify-center mt-6'):
            def refresh_all():
                refresh_stats()
                # Refresh current tab based on which one is active
                load_recent_activity()
                load_errors()
                load_user_actions()
                load_auth_logs()
                load_system_events()

                ui.notify('📊 Dashboard refreshed!', type='positive')

            ui.button('🔄 Refresh Dashboard', on_click=refresh_all).props('color=primary size=lg').classes('px-8 py-3')


# Helper functions for different card types:

def display_clean_log_card(log):
    """Display a clean, simplified log card"""
    level = log.get('level', 'INFO')
    timestamp = log.get('timestamp', '')

    # Parse timestamp to show relative time
    try:
        log_time = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
        now = datetime.now()
        diff = now - log_time.replace(tzinfo=None)

        if diff.total_seconds() < 60:
            time_str = 'Just now'
        elif diff.total_seconds() < 3600:
            time_str = f'{int(diff.total_seconds() / 60)}m ago'
        elif diff.total_seconds() < 86400:
            time_str = f'{int(diff.total_seconds() / 3600)}h ago'
        else:
            time_str = log_time.strftime('%m/%d %H:%M')
    except:
        time_str = timestamp[:16] if timestamp else 'Unknown'

    level_colors = {
        'ERROR': 'border-l-red-500 bg-red-50',
        'WARNING': 'border-l-yellow-500 bg-yellow-50',
        'INFO': 'border-l-blue-500 bg-blue-50'
    }
    level_icons = {
        'ERROR': '🚨',
        'WARNING': '⚠️',
        'INFO': 'ℹ️'
    }

    card_class = level_colors.get(level, 'border-l-gray-500 bg-gray-50')
    icon = level_icons.get(level, '📝')

    with ui.card().classes(f'w-full p-3 {card_class} border-l-4'):
        with ui.row().classes('w-full items-center justify-between'):
            with ui.row().classes('items-center gap-3 flex-1'):
                ui.label(icon).classes('text-lg')
                with ui.column().classes('flex-1'):
                    ui.label(log.get('message', 'No message')).classes('text-sm font-medium text-gray-800')
                    if log.get('user_action'):
                        ui.label(f"👤 {log['user_action']}").classes('text-xs text-blue-600')

            with ui.column().classes('text-right'):
                ui.label(time_str).classes('text-xs text-gray-500')
                ui.label(f"[{log.get('component', 'system')}]").classes('text-xs text-gray-400')


def display_error_card(log):
    """Display detailed error card"""
    with ui.card().classes('w-full p-4 border-l-4 border-l-red-500 bg-red-50'):
        with ui.row().classes('w-full items-start gap-3'):
            ui.icon('error').classes('text-2xl text-red-600 mt-1')
            with ui.column().classes('flex-1'):
                ui.label(log.get('message', 'Error occurred')).classes('text-base font-semibold text-red-800')
                ui.label(log.get('timestamp', '')).classes('text-xs text-red-600 mb-2')

                if log.get('error_details'):
                    ui.label(f"Details: {log['error_details']}").classes('text-sm text-red-700 bg-red-100 p-2 rounded')

                if log.get('user_action'):
                    ui.label(f"User was: {log['user_action']}").classes('text-xs text-red-600 mt-1')


def display_user_action_card(log):
    """Display user interaction card"""
    with ui.card().classes('w-full p-3 border-l-4 border-l-green-500 bg-green-50'):
        with ui.row().classes('w-full items-center gap-3'):
            ui.icon('person').classes('text-xl text-green-600')
            with ui.column().classes('flex-1'):
                action = log.get('user_action', 'User action')
                ui.label(action.replace('_', ' ').title()).classes('text-sm font-medium text-green-800')
                ui.label(log.get('timestamp', '')[:16]).classes('text-xs text-green-600')

            if 'upload' in log.get('message', '').lower():
                ui.icon('upload_file').classes('text-lg text-green-600')


def display_auth_card(log):
    """Display authentication event card"""
    is_success = 'successful' in log.get('message', '').lower()
    color = 'green' if is_success else 'red'
    icon = '✅' if is_success else '❌'

    with ui.card().classes(f'w-full p-3 border-l-4 border-l-{color}-500 bg-{color}-50'):
        with ui.row().classes('w-full items-center gap-3'):
            ui.label(icon).classes('text-lg')
            with ui.column().classes('flex-1'):
                ui.label(log.get('message', 'Auth event')).classes(f'text-sm font-medium text-{color}-800')
                ui.label(log.get('timestamp', '')[:16]).classes(f'text-xs text-{color}-600')


def display_system_card(log):
    """Display system event card"""
    with ui.card().classes('w-full p-3 border-l-4 border-l-blue-500 bg-blue-50'):
        with ui.row().classes('w-full items-center gap-3'):
            ui.icon('settings').classes('text-lg text-blue-600')
            with ui.column().classes('flex-1'):
                ui.label(log.get('message', 'System event')).classes('text-sm font-medium text-blue-800')
                ui.label(log.get('timestamp', '')[:16]).classes('text-xs text-blue-600')



@ui.page('/')
def main_page():
    add_security_headers()  # MOVE THIS TO THE TOP
    ui.add_head_html('<link rel="icon" href="assets/favicon.ico">')
    ui.add_head_html("""
        <script src="https://unpkg.com/@lottiefiles/lottie-player@latest/dist/lottie-player.js"></script>
        """)

    session = app.storage.user
    session.uploaded_file_path = None
    session.uploaded_file_name = "Notes"

    def session_keeper():
        while True:
            try:
                time.sleep(10)  # Ping every 10 seconds
                keep_session_alive()
            except:
                break

    import threading
    threading.Thread(target=session_keeper, daemon=True).start()

    def show_user_error(error_type: str, additional_context: str = ""):
        """Display user-friendly error messages"""
        title, message = ErrorHandler.get_user_friendly_message(error_type)

        if additional_context:
            message += f"\n\n{additional_context}"

        error_label.text = f"{title}\n{message}"
        error_label.visible = True
        try_again_button.visible = True
        error_report_button.visible = True

        # Hide all processing animations
        text_extraction_animation.visible = False
        notes_generation_animation.visible = False
        word_file_generation_animation.visible = False
        generate_button.visible = False
        status_label.text = ""

    def reset_app():
        """Reset application state"""
        try:
            logger.info("Resetting application state")
            session.uploaded_file_path = None
            session.uploaded_file_name = "Notes"
            status_label.text = ""
            download_button.visible = False
            generate_button.visible = True
            reset_button.visible = False
            error_label.text = ""
            error_label.visible = False
            try_again_button.visible = False
            feedback_button.visible = False
            error_report_button.visible = False
            render_upload()
            ui.notify("Ready for another file!")
        except Exception as e:
            ErrorHandler.log_error("RESET_ERROR", str(e), "Resetting application")

    async def process_with_ai():
        """Main processing function with comprehensive error handling"""
        if not session.uploaded_file_path or not session.uploaded_file_path.exists():
            ui.notify(message='Please Upload a File', type='warning')
            return

        ui.notify('Processing may take 5-15 minutes. Mobile devices may experience connection issues.', type='info',
                  timeout=5000)

        async def background_job():
            try:
                logger.info(f"Starting AI processing for file: {session.uploaded_file_path}")

                generate_button.visible = False
                text_extraction_animation.visible = True
                status_label.text = "Extracting content from the document..."
                error_label.visible = False
                ui.update()

                # TEXT EXTRACTION PHASE
                try:
                    logger.info("Starting text extraction phase")
                    extracted_json = await run.io_bound(
                        lambda: send_msg_to_ai(session.uploaded_file_path)
                    )

                    # Check for extraction errors
                    if isinstance(extracted_json, str):
                        if "API" in extracted_json.upper() and "KEY" in extracted_json.upper():
                            ErrorHandler.log_error("API_KEY_MISSING", extracted_json, "Text extraction")
                            show_user_error("API_KEY_MISSING")
                        elif "QUOTA" in extracted_json.upper() or "LIMIT" in extracted_json.upper():
                            ErrorHandler.log_error("API_QUOTA_EXCEEDED", extracted_json, "Text extraction")
                            show_user_error("API_QUOTA_EXCEEDED")
                        elif "AUTHENTICATION" in extracted_json.upper() or "AUTH" in extracted_json.upper():
                            ErrorHandler.log_error("API_AUTHENTICATION_ERROR", extracted_json, "Text extraction")
                            show_user_error("API_AUTHENTICATION_ERROR")
                        elif "NETWORK" in extracted_json.upper() or "CONNECTION" in extracted_json.upper():
                            ErrorHandler.log_error("NETWORK_ERROR", extracted_json, "Text extraction")
                            show_user_error("NETWORK_ERROR")
                        else:
                            ErrorHandler.log_error("EXTRACTION_FAILED", extracted_json, "Text extraction")
                            show_user_error("EXTRACTION_FAILED")
                        return

                    if not extracted_json or len(extracted_json) == 0:
                        ErrorHandler.log_error("EXTRACTION_EMPTY", "No content extracted", "Text extraction")
                        show_user_error("EXTRACTION_FAILED", "The document appears to be empty or unreadable.")
                        return

                    logger.info("Text extraction completed successfully")

                except Exception as e:
                    error_msg = str(e)
                    if "timeout" in error_msg.lower():
                        ErrorHandler.log_error("PROCESSING_TIMEOUT", error_msg, "Text extraction")
                        show_user_error("PROCESSING_TIMEOUT")
                    elif "network" in error_msg.lower() or "connection" in error_msg.lower():
                        ErrorHandler.log_error("NETWORK_ERROR", error_msg, "Text extraction")
                        show_user_error("NETWORK_ERROR")
                    else:
                        ErrorHandler.log_error("EXTRACTION_UNEXPECTED_ERROR", error_msg, "Text extraction")
                        show_user_error("EXTRACTION_FAILED")
                    return

                # NOTES GENERATION PHASE
                text_extraction_animation.visible = False
                status_label.text = "🛠 Generating Notes"
                notes_generation_animation.visible = True
                ui.update()

                try:
                    logger.info("Starting notes generation phase")
                    notes_generated = await run.io_bound(generate_notes_from_content, extracted_json)

                    if not notes_generated:
                        ErrorHandler.log_error("NOTES_GENERATION_EMPTY", "No notes generated", "Notes generation")
                        show_user_error("NOTES_GENERATION_FAILED")
                        return

                    if isinstance(notes_generated, str) and "error" in notes_generated.lower():
                        if "api" in notes_generated.lower() and "key" in notes_generated.lower():
                            ErrorHandler.log_error("API_KEY_MISSING", notes_generated, "Notes generation")
                            show_user_error("API_KEY_MISSING")
                        elif "quota" in notes_generated.lower() or "limit" in notes_generated.lower():
                            ErrorHandler.log_error("API_QUOTA_EXCEEDED", notes_generated, "Notes generation")
                            show_user_error("API_QUOTA_EXCEEDED")
                        else:
                            ErrorHandler.log_error("NOTES_GENERATION_ERROR", notes_generated, "Notes generation")
                            show_user_error("NOTES_GENERATION_FAILED")
                        return

                    logger.info("Notes generation completed successfully")

                except Exception as e:
                    error_msg = str(e)
                    if "timeout" in error_msg.lower():
                        ErrorHandler.log_error("PROCESSING_TIMEOUT", error_msg, "Notes generation")
                        show_user_error("PROCESSING_TIMEOUT")
                    elif "quota" in error_msg.lower() or "limit" in error_msg.lower():
                        ErrorHandler.log_error("API_QUOTA_EXCEEDED", error_msg, "Notes generation")
                        show_user_error("API_QUOTA_EXCEEDED")
                    else:
                        ErrorHandler.log_error("NOTES_GENERATION_UNEXPECTED_ERROR", error_msg, "Notes generation")
                        show_user_error("NOTES_GENERATION_FAILED")
                    return

                # WORD FILE GENERATION PHASE
                notes_generation_animation.visible = False
                status_label.text = "📕 Preparing Word file..."
                word_file_generation_animation.visible = True
                ui.update()

                try:
                    logger.info("Starting Word file generation phase")
                    unique_name = f"{session.uploaded_file_name}_{uuid.uuid4().hex[:6]}.docx"
                    file_generated = await run.io_bound(generate_word_file, notes_generated,
                                                        file_name=unique_name.replace(' ', '_'))

                    if not file_generated or not os.path.exists(file_generated):
                        ErrorHandler.log_error("WORD_FILE_NOT_CREATED", "File generation returned empty path",
                                               "Word file generation")
                        show_user_error("WORD_FILE_GENERATION_FAILED")
                        return

                    # Prepare download
                    with open(file_generated, 'rb') as f:
                        file_content = f.read()
                    os.remove(file_generated)

                    base64_data = base64.b64encode(file_content).decode('utf-8')
                    mime_type = mimetypes.guess_type("Notes.docx")[0] or "application/octet-stream"

                    def trigger_download():
                        ui.run_javascript(f"""
                            const link = document.createElement('a');
                            link.href = "data:{mime_type};base64,{base64_data}";
                            link.download = "{session.uploaded_file_name}_Notes.docx";
                            link.click();
                        """)

                    time.sleep(3)

                    download_button.on('click', trigger_download, [])
                    word_file_generation_animation.visible = False
                    download_button.visible = True
                    feedback_button.visible = True
                    status_label.text = "✅ Your Notes are Ready!"
                    reset_button.visible = True
                    ui.update()

                    logger.info("Processing completed successfully")

                except Exception as e:
                    ErrorHandler.log_error("WORD_FILE_GENERATION_ERROR", str(e), "Word file generation")
                    show_user_error("WORD_FILE_GENERATION_FAILED")
                    return

            except Exception as e:
                # Catch any unexpected failures
                ErrorHandler.log_error("PROCESSING_UNEXPECTED_ERROR", str(e), "Background processing")
                show_user_error("UNKNOWN_ERROR")

        background_tasks.create(background_job())

    # UI Layout
    with ui.column().classes(
            'absolute inset-0 w-full h-full overflow-x-hidden bg-gradient-to-tr '
            'from-emerald-200 via-white to-indigo-100 px-3 sm:px-6 py-6 sm:py-12 text-base sm:text-lg'):

        with ui.column().classes('w-full items-center'):
            ui.label('NotesCraft AI') \
                .classes('text-4xl md:text-4xl font-extrabold text-emerald-800 text-center')

            ui.label('Transform your PDFs into beautiful, structured study notes.') \
                .classes('text-base md:text-lg text-gray-600 text-center mb-6 px-4')

        with ui.card().classes(
                'w-full max-w-2xl sm:max-w-3xl mx-auto p-6 sm:p-8 bg-white/90 backdrop-blur-lg shadow-2xl rounded-3xl border border-gray-200 flex flex-col items-center space-y-4'):
            upload_container = ui.column().classes('w-full items-center')

            def handle_upload(e):

                ui.run_javascript("""
                       uploadInProgress = false;
                       wasHidden = false;
                       console.log('Upload received successfully');
                   """)

                try:
                    logger.info(f"File upload initiated: {e.name}")

                    temp_file_name = e.name
                    temp_content = e.content.read()

                    if not temp_content:
                        logger.warning("Empty file uploaded")
                        ui.notify("⚠️ Upload failed: empty file received.", type="warning")
                        return

                    file_size = len(temp_content)
                    suffix = Path(temp_file_name).suffix or ".pdf"

                    # Create temporary file for validation
                    with NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                        temp_file.write(temp_content)
                        temp_file.flush()
                        os.fsync(temp_file.fileno())
                        temp_file_path = Path(temp_file.name)

                    # Validate file
                    is_valid, error_type, page_count = validate_file(temp_file_path, file_size)

                    if not is_valid:
                        os.unlink(temp_file_path)
                        title, message = ErrorHandler.get_user_friendly_message(error_type)
                        ui.notify(f"{title}: {message}", type="negative", timeout=8000)
                        ErrorHandler.log_error("FILE_VALIDATION_FAILED", f"{error_type} for file {temp_file_name}",
                                               "File upload")
                        return

                    def confirm_upload():
                        session.uploaded_file_name = temp_file_name
                        session.uploaded_file_path = temp_file_path
                        dialog.close()

                        upload_container.clear()
                        with upload_container:
                            with ui.card().classes(
                                    'w-full max-w-xl p-6 sm:p-8 rounded-2xl border border-emerald-300 '
                                    'bg-emerald-50 text-center shadow-md flex flex-col items-center justify-center'
                            ):
                                ui.icon("picture_as_pdf").classes("text-red-500 text-5xl sm:text-6xl")
                                ui.label(session.uploaded_file_name).classes(
                                    "text-lg sm:text-xl font-semibold text-gray-800 mt-2")
                                ui.label(f"✅ File Uploaded Successfully ({page_count} pages)").classes(
                                    "text-sm sm:text-base text-gray-600 mt-1")

                        logger.info(f"File upload confirmed: {temp_file_name}")

                    def cancel_upload():
                        os.unlink(temp_file_path)
                        dialog.close()
                        ui.notify("Upload cancelled", type="info")
                        logger.info(f"File upload cancelled: {temp_file_name}")

                    # Show confirmation dialog
                    with ui.dialog() as dialog, ui.card().classes('w-full max-w-md p-6'):
                        ui.label('Confirm File Upload').classes('text-xl font-bold text-gray-800 mb-4')

                        with ui.column().classes('w-full items-center space-y-3'):
                            ui.icon("description").classes("text-blue-500 text-4xl")
                            ui.label(f'File: {temp_file_name}').classes('text-lg font-medium text-gray-700')
                            ui.label(f'Pages: {page_count}').classes('text-base text-gray-600')
                            ui.label(f'Size: {file_size / 1024 / 1024:.1f} MB').classes('text-base text-gray-600')

                            ui.label('Do you want to upload this file?').classes(
                                'text-base text-gray-700 mt-4 text-center')

                            with ui.row().classes('w-full justify-center gap-4 mt-6'):
                                ui.button('Yes, Upload', on_click=confirm_upload).props(
                                    'unelevated rounded color=green text-color=white').classes(
                                    'px-6 py-2 font-semibold')
                                ui.button('Cancel', on_click=cancel_upload).props(
                                    'unelevated rounded color=red text-color=white').classes(
                                    'px-6 py-2 font-semibold')

                    dialog.open()

                except Exception as e:
                    ErrorHandler.log_error("UPLOAD_HANDLER_ERROR", str(e), "File upload handling")
                    ui.notify("Upload failed due to an unexpected error. Please try again.", type="negative")

            def render_upload():
                upload_container.clear()
                with upload_container:
                    uploader = ui.upload(label='', on_upload=handle_upload, auto_upload=True, multiple=False).props(
                        'accept=.pdf,.docx').classes('hidden')

                    with ui.card().classes(
                            'w-full max-w-xl h-48 border-2 border-dashed border-gray-300 bg-white/80 '
                            'hover:bg-emerald-50 rounded-2xl flex flex-col items-center justify-center '
                            'cursor-pointer transition-all text-center shadow-md'
                    ).on('click', lambda: uploader.run_method('pickFiles')):
                        ui.icon('cloud_upload').classes('text-5xl text-emerald-600')
                        ui.label('Click to upload your PDF or DOCX').classes('text-lg font-medium text-gray-700')
                        ui.label('or drag and drop here').classes('text-sm text-gray-500')

                    ui.label(f'Maximum {MAX_PAGES} pages allowed').classes('text-xs text-gray-400 mt-2 text-center')

            render_upload()

            # Processing animations and status
            with ui.column().classes('w-full items-center'):
                with ui.column().classes('w-full items-center sm:flex-row sm:justify-center sm:gap-6 mt-2'):
                    text_extraction_animation = ui.html("""
                        <lottie-player src="/assets/document-search.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    text_extraction_animation.visible = False

                    notes_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_notes.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    notes_generation_animation.visible = False

                    word_file_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_word_file.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    word_file_generation_animation.visible = False

                status_label = ui.label('').classes(
                    'text-gray-900 mt-3 text-center text-base sm:text-lg font-medium'
                )

                error_label = ui.label('').classes(
                    'mt-3 text-base sm:text-lg text-rose-500 font-semibold text-center whitespace-pre-line')
                error_label.visible = False

            # Action buttons
            download_button = ui.button('Download Notes').props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            download_button.visible = False

            feedback_url = "https://forms.gle/gPHd66XpZ1nM17si9"

            def open_feedback_form():
                js_code = f"window.open('{feedback_url}', '_blank');"
                ui.run_javascript(js_code)

            feedback_button = ui.button("📝 Give Feedback", on_click=open_feedback_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            feedback_button.visible = False

            error_report_url = "https://forms.gle/Eqjk6SS1jtmWuXGg6"

            def open_error_report_form():
                js_code = f"window.open('{error_report_url}', '_blank');"
                ui.run_javascript(js_code)

            error_report_button = ui.button("🚩 Report Error", on_click=open_error_report_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            error_report_button.visible = False

            generate_button = ui.button('🚀 Generate Notes', on_click=process_with_ai).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')

            reset_button = ui.button('🔄 Upload Another File', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            reset_button.visible = False

            try_again_button = ui.button('🔄 Try Again', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            try_again_button.visible = False


def validate_environment():
    """Validate critical environment variables"""

    # Check if LOG_ADMIN_PASSWORD is set in production
    if not os.environ.get('LOG_ADMIN_PASSWORD'):
        if os.environ.get('RAILWAY_ENVIRONMENT') or os.environ.get('HEROKU_APP_NAME'):
            log_warning("security", "LOG_ADMIN_PASSWORD not set in production environment",
                        user_action="startup_check",
                        additional_data={"env": "production"})
            print("⚠️  WARNING: LOG_ADMIN_PASSWORD environment variable not set in production!")
        else:
            log_info("security", "Development mode: Using generated password",
                     user_action="startup_check",
                     additional_data={"env": "development"})


validate_environment()


ui.run(
    host='0.0.0.0',
    port=int(os.environ.get('PORT', 8080)),
    storage_secret=os.environ.get('STORAGE_SECRET', secrets.token_hex(32)),
    title='NotesCraft AI – Smart Notes Maker',
    # ADD these security options:
    reload=False,  # Disable auto-reload in production
    show=False,    # Don't auto-open browser
    favicon='assets/favicon.ico'
)