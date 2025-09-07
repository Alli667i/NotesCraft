import base64
import mimetypes
import os
import time
import uuid
from pathlib import Path
from tempfile import NamedTemporaryFile
import secrets
import requests
from nicegui import ui, app, background_tasks, run
from process_content_to_notest_test_fix import generate_notes_from_content
from process_pdf_to_json_test import send_msg_to_ai
from process_to_word_02 import generate_word_file

# Import libraries for page counting
import fitz  # PyMuPDF for PDFs
from docx import Document  # python-docx for DOCX files

app.add_static_files('/assets', os.path.join(os.path.dirname(__file__), 'assets'))

# Ensure root container fills full viewport on mobile
ui.add_head_html("""
<style>
html, body, #__nicegui_root {
    height: 100%;
    margin: 0;
    padding: 0;
}
</style>
""")

# File validation constants
MAX_PAGES = 20
MAX_FILE_SIZE_MB = 70  # Additional safety check


def count_pages(file_path):
    """
    Count pages in PDF or DOCX files
    Returns: (page_count, error_message)
    """
    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            # Count PDF pages using PyMuPDF
            doc = fitz.open(file_path)
            page_count = doc.page_count
            doc.close()
            return page_count, None

        elif file_extension == '.docx':
            # Count DOCX pages (approximate based on content)
            doc = Document(file_path)

            # Method 1: Count paragraphs and estimate pages
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])

            # Rough estimation: ~10-15 paragraphs per page
            estimated_pages = max(1, paragraph_count // 12)

            # Method 2: Count characters and estimate (more accurate)
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            char_based_pages = max(1, total_chars // 3000)  # ~3000 chars per page

            # Use the higher estimate for safety
            page_count = max(estimated_pages, char_based_pages)

            return page_count, None

        else:
            return 0, "Unsupported file format"

    except Exception as e:
        return 0, f"Error reading file: {str(e)}"


def validate_file(file_path, file_size_bytes):
    """
    Validate uploaded file against size and page limits
    Returns: (is_valid, error_message, page_count)
    """
    # Check file size first (quick check)
    file_size_mb = file_size_bytes / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        return False, f"File too large ({file_size_mb:.1f}MB). Maximum size is {MAX_FILE_SIZE_MB}MB.", 0

    # Count pages
    page_count, error = count_pages(file_path)

    if error:
        return False, error, 0

    if page_count > MAX_PAGES:
        return False, f"Document has {page_count} pages. Maximum allowed is {MAX_PAGES} pages.", page_count

    return True, None, page_count


@ui.page('/')
def main_page():
    ui.add_head_html('<link rel="icon" href="assets/favicon.ico">')
    ui.add_head_html("""
    <script src="https://unpkg.com/@lottiefiles/lottie-player@latest/dist/lottie-player.js"></script>
    """)

    session = app.storage.user
    session.uploaded_file_path = None
    session.uploaded_file_name = "Notes"

    # --- Helper Functions ---
    def report_error(Error):
        if Error:
            try:
                requests.post(
                    'https://script.google.com/macros/s/AKfycbz6Gbht0iZ4tW7lp48x3hDYCvYIDGZbOYdwnpbmyHSQjxsdZ0D0zsx7ZU84eN9n0g2T9w/exec',
                    json={"Error": Error},
                )
            except Exception as e:
                print(f"Error reporting failed: {e}")

    def reset_app():
        session.uploaded_file_path = None
        session.uploaded_file_name = "Notes"
        status_label.text = ""
        download_button.visible = False
        generate_button.visible = True
        reset_button.visible = False
        error_label.text = ""
        try_again_button.visible = False
        feedback_button.visible = False
        error_report_button.visible = False
        render_upload()
        ui.notify("Ready for another file!")

    async def process_with_ai():
        if not session.uploaded_file_path or not session.uploaded_file_path.exists():
            ui.notify(message='Please Upload a File', type='warning')
            return

        ui.notify('Processing may take 5-10 minutes. Mobile devices may experience connection issues.', type='info',
                  timeout=5000)

        async def background_job():
            try:
                generate_button.visible = False
                text_extraction_animation.visible = True
                status_label.text = "Extracting content from the document..."
                ui.update()

                # --- TEXT EXTRACTION ---
                try:
                    extracted_json = await run.io_bound(
                        lambda: send_msg_to_ai(session.uploaded_file_path, report_error)
                    )

                    # ✅ Check if extraction failed with an error string
                    if isinstance(extracted_json, str) and extracted_json.startswith("Extraction Failed"):
                        error_report_button.visible = True
                        text_extraction_animation.visible = False
                        try_again_button.visible = True
                        status_label.text = ""
                        error_label.text = f"⚠️ {extracted_json}"
                        ui.update()
                        return

                except Exception as e:
                    report_error(f"Unexpected Background Error: {str(e)}")
                    error_report_button.visible = True
                    text_extraction_animation.visible = False
                    try_again_button.visible = True
                    error_label.text = f"⚠️ Unexpected Error: {str(e)}"
                    ui.update()
                    return

                # --- NOTES GENERATION ---
                text_extraction_animation.visible = False
                status_label.text = "🛠 Generating Notes"
                notes_generation_animation.visible = True
                ui.update()

                try:
                    notes_generated = await run.io_bound(generate_notes_from_content, extracted_json)
                    if not notes_generated:
                        raise ValueError("Notes generation returned empty content.")
                except Exception as e:
                    report_error(f"Notes Generation Error: {str(e)}")
                    print(f"Error: {str(e)}")
                    error_report_button.visible = True
                    notes_generation_animation.visible = False
                    status_label.text = ""
                    error_label.text = "⚠️ Failed to generate notes. Please try again."
                    try_again_button.visible = True
                    ui.update()
                    return

                # --- WORD FILE CREATION ---
                notes_generation_animation.visible = False
                status_label.text = "📕 Preparing Word file..."
                word_file_generation_animation.visible = True
                ui.update()

                unique_name = f"{session.uploaded_file_name}_{uuid.uuid4().hex[:6]}.docx"
                file_generated = await run.io_bound(generate_word_file, notes_generated,
                                                    file_name=unique_name.replace(' ', '_'))

                # --- PREPARE DOWNLOAD ---
                with open(file_generated, 'rb') as f:
                    file_content = f.read()
                os.remove(file_generated)

                base64_data = base64.b64encode(file_content).decode('utf-8')
                mime_type = mimetypes.guess_type("Notes.docx")[0] or "application/octet-stream"

                def trigger_download():
                    ui.run_javascript(f"""
                        const link = document.createElement('a');
                        link.href = "data:{mime_type};base64,{base64_data}";
                        link.download = "{session.uploaded_file_name}_Notes.docx";
                        link.click();
                    """)

                time.sleep(3)

                # Make sure only one listener is active
                download_button.on('click', trigger_download, [])
                word_file_generation_animation.visible = False
                download_button.visible = True
                feedback_button.visible = True
                status_label.text = "✅ Your Notes are Ready!"
                reset_button.visible = True
                ui.update()

            except Exception as e:
                # --- Catch any unexpected failures ---
                report_error(str(e))
                error_report_button.visible = True
                print(f"Error: {str(e)}")
                error_label.text = "⚠️ Something went wrong during processing. Please try again."
                status_label.text = ""
                text_extraction_animation.visible = False
                notes_generation_animation.visible = False
                word_file_generation_animation.visible = False
                try_again_button.visible = True
                ui.update()

        background_tasks.create(background_job())

    # --- UI Layout ---
    with ui.column().classes(
            'absolute inset-0 w-full h-full overflow-x-hidden bg-gradient-to-tr '
            'from-emerald-200 via-white to-indigo-100 px-3 sm:px-6 py-6 sm:py-12 text-base sm:text-lg'):

        with ui.column().classes('w-full items-center'):
            ui.label('NotesCraft AI') \
                .classes('text-4xl md:text-4xl font-extrabold text-emerald-800 text-center')

            ui.label('Transform your PDFs into beautiful, structured study notes.') \
                .classes('text-base md:text-lg text-gray-600 text-center mb-6 px-4')

        with ui.card().classes(
                'w-full max-w-2xl sm:max-w-3xl mx-auto p-6 sm:p-8 bg-white/90 backdrop-blur-lg shadow-2xl rounded-3xl border border-gray-200 flex flex-col items-center space-y-4'):
            upload_container = ui.column().classes('w-full items-center')

            def handle_upload(e):
                # Store file details temporarily for confirmation
                temp_file_name = e.name
                temp_content = e.content.read()

                if not temp_content:
                    ui.notify("⚠️ Upload failed: empty file received.", type="warning")
                    return

                file_size = len(temp_content)
                suffix = Path(temp_file_name).suffix or ".pdf"

                # Create temporary file for validation
                with NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                    temp_file.write(temp_content)
                    temp_file.flush()
                    os.fsync(temp_file.fileno())
                    temp_file_path = Path(temp_file.name)

                # Validate file before showing confirmation
                is_valid, error_msg, page_count = validate_file(temp_file_path, file_size)

                if not is_valid:
                    # Clean up temp file and show error
                    os.unlink(temp_file_path)
                    ui.notify(f"⚠️ {error_msg}", type="negative", timeout=5000)
                    return

                def confirm_upload():
                    # User confirmed - proceed with upload
                    session.uploaded_file_name = temp_file_name
                    session.uploaded_file_path = temp_file_path

                    dialog.close()

                    # Show success UI with page count
                    upload_container.clear()
                    with upload_container:
                        with ui.card().classes(
                                'w-full max-w-xl p-6 sm:p-8 rounded-2xl border border-emerald-300 '
                                'bg-emerald-50 text-center shadow-md flex flex-col items-center justify-center'
                        ):
                            ui.icon("picture_as_pdf").classes("text-red-500 text-5xl sm:text-6xl")
                            ui.label(session.uploaded_file_name).classes(
                                "text-lg sm:text-xl font-semibold text-gray-800 mt-2")
                            ui.label(f"✅ File Uploaded Successfully ({page_count} pages)").classes(
                                "text-sm sm:text-base text-gray-600 mt-1")

                def cancel_upload():
                    # User cancelled - clean up temp file
                    os.unlink(temp_file_path)
                    dialog.close()
                    ui.notify("Upload cancelled", type="info")

                # Show confirmation dialog
                with ui.dialog() as dialog, ui.card().classes('w-full max-w-md p-6'):
                    ui.label('Confirm File Upload').classes('text-xl font-bold text-gray-800 mb-4')

                    with ui.column().classes('w-full items-center space-y-3'):
                        ui.icon("description").classes("text-blue-500 text-4xl")
                        ui.label(f'File: {temp_file_name}').classes('text-lg font-medium text-gray-700')
                        ui.label(f'Pages: {page_count}').classes('text-base text-gray-600')
                        ui.label(f'Size: {file_size / 1024 / 1024:.1f} MB').classes('text-base text-gray-600')

                        ui.label('Do you want to upload this file?').classes('text-base text-gray-700 mt-4 text-center')

                        with ui.row().classes('w-full justify-center gap-4 mt-6'):
                            ui.button('Yes, Upload', on_click=confirm_upload).props(
                                'unelevated rounded color=green text-color=white').classes(
                                'px-6 py-2 font-semibold')
                            ui.button('Cancel', on_click=cancel_upload).props(
                                'unelevated rounded color=red text-color=white').classes(
                                'px-6 py-2 font-semibold')

                dialog.open()

            def render_upload():
                upload_container.clear()
                with upload_container:
                    uploader = ui.upload(label='', on_upload=handle_upload, auto_upload=True, multiple=False).props(
                        'accept=.pdf,.docx').classes('hidden')

                    with ui.card().classes(
                            'w-full max-w-xl h-48 border-2 border-dashed border-gray-300 bg-white/80 '
                            'hover:bg-emerald-50 rounded-2xl flex flex-col items-center justify-center '
                            'cursor-pointer transition-all text-center shadow-md'
                    ).on('click', lambda: uploader.run_method('pickFiles')):
                        ui.icon('cloud_upload').classes('text-5xl text-emerald-600')
                        ui.label('Click to upload your PDF or DOCX').classes('text-lg font-medium text-gray-700')
                        ui.label('or drag and drop here').classes('text-sm text-gray-500')

                    # Add page limit info below the upload area
                    ui.label(f'Maximum {MAX_PAGES} pages allowed').classes('text-xs text-gray-400 mt-2 text-center')

            render_upload()

            with ui.column().classes('w-full items-center'):
                with ui.column().classes('w-full items-center sm:flex-row sm:justify-center sm:gap-6 mt-2'):
                    text_extraction_animation = ui.html("""
                        <lottie-player src="/assets/document-search.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    text_extraction_animation.visible = False

                    notes_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_notes.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    notes_generation_animation.visible = False

                    word_file_generation_animation = ui.html("""
                        <lottie-player src="/assets/generate_word_file.json" background="transparent" speed="1"
                                       style="width: 130px; height: 130px;" loop autoplay></lottie-player>
                    """)
                    word_file_generation_animation.visible = False

                status_label = ui.label('').classes(
                    'text-gray-900 mt-3 text-center text-base sm:text-lg font-medium'
                )

                error_label = ui.label('').classes('mt-3 text-base sm:text-lg text-rose-500 font-semibold text-center')

            download_button = ui.button('Download Notes').props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            download_button.visible = False

            feedback_url = "https://forms.gle/gPHd66XpZ1nM17si9"

            def open_feedback_form():
                js_code = f"window.open('{feedback_url}', '_blank');"
                ui.run_javascript(js_code)

            feedback_button = ui.button("📝 Give Feedback", on_click=open_feedback_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            feedback_button.visible = False

            error_report_url = "https://forms.gle/Eqjk6SS1jtmWuXGg6"

            def open_error_report_form():
                js_code = f"window.open('{error_report_url}', '_blank');"
                ui.run_javascript(js_code)

            error_report_button = ui.button("🚩 Report Error ", on_click=open_error_report_form).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            error_report_button.visible = False

            generate_button = ui.button('🚀 Generate Notes', on_click=process_with_ai).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')

            reset_button = ui.button('🔄 Upload Another File', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            reset_button.visible = False

            try_again_button = ui.button('🔄 Try Again', on_click=reset_app).props(
                'unelevated rounded color=indigo text-color=white').classes(
                'w-full max-w-md mx-auto mt-4 sm:mt-6 px-4 sm:px-6 py-3 text-base sm:text-lg font-semibold shadow-sm transition-all duration-200 text-center')
            try_again_button.visible = False


ui.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)),
       storage_secret=os.environ.get('STORAGE_SECRET', secrets.token_hex(32)),
       title='NotesCraft AI – Smart Notes Maker')