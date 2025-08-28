from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_formatted_text(paragraph, text, font_size=12):
    """
    Applies bold for **...** segments while keeping the rest normal.
    Example: 'This is **bold** and normal' -> correct runs with bold.
    """
    # Split into bold/non-bold parts; non-greedy to support multiple bold segments
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])  # strip ** **
            run.bold = True
        else:
            run = paragraph.add_run(part)
        # Ensure consistent font on every run
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)

def add_paragraph_with_spacing(doc, text, font_size=12, bold=False, align='left',
                               space_before=12, space_after=6):
    paragraph = doc.add_paragraph()
    # Start with an optional bold run if the whole line should be bold (e.g., headings),
    # but still respect inline **bold** segments.
    if bold:
        # We’ll build from formatted text; bold=True means default style is bold unless overridden by **...**
        # Easiest: apply formatted runs, then set paragraph runs to bold if not already bolded.
        add_formatted_text(paragraph, text, font_size)
        for run in paragraph.runs:
            if run.bold is None:  # only set if not already explicitly bold
                run.bold = True
    else:
        add_formatted_text(paragraph, text, font_size)

    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph

def generate_word_file(content, file_name):
    doc = Document()

    # Optional: safely set Normal style without deprecated lookup
    # (Iterate styles and match by name & type)
    normal_style = None
    for s in doc.styles:
        if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name.lower() == 'normal':
            normal_style = s
            break
    if normal_style:
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)

    for item in content:
        try:
            item_type = item.get('type')
            text = item.get('text', '') or ''

            if item_type == 'heading':
                add_paragraph_with_spacing(
                    doc, text, font_size=16, bold=True,
                    align='center', space_before=48, space_after=36
                )

            elif item_type == 'subheading':
                add_paragraph_with_spacing(
                    doc, text, font_size=14, bold=True,
                    align='left', space_before=36, space_after=24
                )

            elif item_type == 'paragraph':
                add_paragraph_with_spacing(
                    doc, text, font_size=12, bold=False,
                    align='justify', space_before=24, space_after=24
                )



            elif item_type == 'bullet':

                bullet_para = doc.add_paragraph()

                add_formatted_text(bullet_para, f"• {text}", font_size=12)

                pf = bullet_para.paragraph_format

                pf.space_before = Pt(24)

                pf.space_after = Pt(28)

                pf.left_indent = Pt(18)

                pf.line_spacing = Pt(16)  # forces a "big gap" look


        except Exception as e:
            print(f"Skipped invalid item: {item} due to error: {e}")

    file_path = f"{file_name}.docx"
    doc.save(file_path)
    return file_path
